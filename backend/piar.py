"""
Endpoints REST del generador de PIAR (Plan Individual de Ajustes Razonables).

Contrato:
- POST /api/piar/                      → crear borrador desde conversación (síntesis vía Claude)
- GET  /api/piar/estudiante/{eid}      → listar todas las versiones del estudiante
- PUT  /api/piar/{piar_id}/aprobar     → aprobar (borrador → aprobado, inmutable)
- GET  /api/piar/{piar_id}/docx        → descargar DOCX on-demand (marca BORRADOR si aplica)

Todos los endpoints validan que el PIAR pertenezca al docente autenticado.

Nota de diseño (aprobada por el owner):
- Denormalización de id_grupo e id_docente en la tabla PIAR (derivables
  desde id_estudiante) — asumida para queries frecuentes sin JOINs.
- Sin docx_bytes: el DOCX se regenera on-demand desde `contenido` cada vez
  que se pide. Permite cambiar el template sin migrar datos.
"""
from __future__ import annotations

import io
import json
import re
from datetime import datetime
from typing import List, Optional

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session

from auth import get_current_docente
from database import get_db
from ia import client as anthropic_client
from config import settings
from models import Docente, Estudiante, Grupo, Mensaje, PIAR
from prompts import MODO_PIAR

router = APIRouter(prefix="/api/piar", tags=["piar"])


# ═══════════════════════════════════════════════════════════════
# SCHEMAS
# ═══════════════════════════════════════════════════════════════

SECCIONES_PIAR = (
    "caracterizacion",
    "barreras",
    "ajustes_razonables",
    "apoyos",
    "metas",
    "seguimiento",
)


class PIARCreateRequest(BaseModel):
    id_estudiante: str
    periodo: int = Field(ge=1, le=4)
    anio: Optional[int] = None  # si no se pasa, se toma grupo.anio_lectivo


class PIAROut(BaseModel):
    id_piar: str
    id_estudiante: str
    id_grupo: str
    id_docente: str
    periodo: int
    anio: int
    version: int
    contenido: dict
    estado: str
    creado_en: datetime
    aprobado_en: Optional[datetime]

    model_config = {"from_attributes": True}


# ═══════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════

def _piar_del_docente_o_404(piar_id: str, docente_id: str, db: Session) -> PIAR:
    p = db.query(PIAR).filter(
        PIAR.id_piar == piar_id,
        PIAR.id_docente == docente_id,
    ).first()
    if not p:
        raise HTTPException(status_code=404, detail="PIAR no encontrado")
    return p


def _next_version(db: Session, id_estudiante: str, id_grupo: str,
                  periodo: int, anio: int) -> int:
    """
    Devuelve version+1 para el tuple (estudiante, grupo, periodo, anio),
    o 1 si no hay versiones previas. Race condition posible pero muy baja
    (un solo docente por grupo). El UNIQUE constraint de la tabla protege
    ante colisiones — en caso de choque, el POST fallará con IntegrityError.
    """
    maxv = (
        db.query(PIAR.version)
        .filter(
            PIAR.id_estudiante == id_estudiante,
            PIAR.id_grupo == id_grupo,
            PIAR.periodo == periodo,
            PIAR.anio == anio,
        )
        .order_by(PIAR.version.desc())
        .limit(1)
        .scalar()
    )
    return (maxv or 0) + 1


def _seccion_pendiente(nombre: str) -> str:
    return f"[PENDIENTE — sin información]"


def _skeleton_pendiente() -> dict:
    """Contenido default cuando la síntesis IA falla o no hay conversación."""
    return {s: _seccion_pendiente(s) for s in SECCIONES_PIAR}


def _sanitizar_contenido(bruto: dict) -> dict:
    """
    Asegura que el JSON de contenido tenga exactamente las 6 secciones
    esperadas. Secciones extra se descartan; secciones ausentes se marcan
    como pendientes. Los valores se coercen a string.
    """
    if not isinstance(bruto, dict):
        return _skeleton_pendiente()
    out = {}
    for s in SECCIONES_PIAR:
        v = bruto.get(s)
        if isinstance(v, str) and v.strip():
            out[s] = v.strip()
        else:
            out[s] = _seccion_pendiente(s)
    return out


async def _sintetizar_conversacion_a_json(
    docente: Docente,
    grupo: Grupo,
    estudiante: Estudiante,
    historial: List[Mensaje],
) -> dict:
    """
    Envía la conversación PIAR completa a Claude con una instrucción
    especial de consolidación y espera que devuelva SOLO el JSON con las
    6 secciones markdown.

    Consume 1 llamada al modelo (contabilizado como uso de rate limit
    fuera de este helper — política aprobada por el owner).

    Si el JSON no parsea, cae al skeleton de pendientes en vez de crashear.
    """
    from ia import _bloque_contexto_grupo, _bloque_piar
    from prompts import PROMPT_BASE, PROMPT_MODO_PIAR

    system_prompt = (
        PROMPT_BASE
        + "\n\n" + PROMPT_MODO_PIAR
        + _bloque_contexto_grupo(grupo, [estudiante])
        + _bloque_piar(estudiante)
    )

    # Turnos de la conversación previa (últimos 40 para no romper contexto)
    messages = []
    for msg in historial[-40:]:
        role = "user" if msg.remitente == "docente" else "assistant"
        if messages and messages[-1]["role"] == role:
            messages[-1]["content"] += "\n\n" + msg.contenido
        else:
            messages.append({"role": role, "content": msg.contenido})

    # Turno final de consolidación
    instruccion_final = (
        "TURNO DE CONSOLIDACIÓN — Sintetizá TODA la conversación anterior en el "
        "JSON estructurado del PIAR con las 6 secciones exactas: "
        f"{', '.join(SECCIONES_PIAR)}. "
        "Cada sección es un string markdown en registro formal, apto para "
        "documento oficial, usando el vocabulario del Decreto 1421 "
        "(BAP, ajustes razonables, apoyos). Las secciones que no se cubrieron "
        "en la conversación marcalas EXACTAMENTE como '[PENDIENTE — sin información]'. "
        "No inventes datos. Devolvé SOLO el JSON, sin texto antes ni después, "
        "sin bloques de código markdown, sin explicaciones."
    )
    if messages and messages[-1]["role"] == "user":
        messages[-1]["content"] += "\n\n" + instruccion_final
    else:
        messages.append({"role": "user", "content": instruccion_final})

    # Llamada sin streaming — respuesta de una sola vez para parseo
    try:
        response = await anthropic_client.messages.create(
            model=settings.CLAUDE_MODEL,
            max_tokens=4096,
            system=system_prompt,
            messages=messages,
        )
        raw = response.content[0].text.strip()
    except Exception as exc:
        # El endpoint captura y decide qué status devolver
        raise RuntimeError(f"Error llamando a Claude para síntesis: {exc}") from exc

    # Robustez: si Claude devuelve el JSON envuelto en ```json ... ```
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)

    try:
        parsed = json.loads(raw)
    except json.JSONDecodeError:
        # No parsea — devolvemos skeleton en vez de romper el flow
        return _skeleton_pendiente()

    return _sanitizar_contenido(parsed)


# ═══════════════════════════════════════════════════════════════
# DOCX GENERATOR (marca BORRADOR si aplica)
# ═══════════════════════════════════════════════════════════════

_SECCION_LABEL = {
    "caracterizacion": "1. Caracterización del estudiante",
    "barreras": "2. Barreras para el aprendizaje y la participación (BAP)",
    "ajustes_razonables": "3. Ajustes razonables",
    "apoyos": "4. Apoyos requeridos",
    "metas": "5. Metas de aprendizaje",
    "seguimiento": "6. Seguimiento y evaluación",
}


def _construir_piar_docx(
    docente: Docente,
    grupo: Grupo,
    estudiante: Estudiante,
    piar: PIAR,
) -> bytes:
    """
    Genera el DOCX del PIAR desde el JSON de contenido. Reutiliza el look
    institucional del generador de documentos IA (backend/documento.py::
    _set_cell_bg y paleta de constantes).

    Encabezado incluye marca "BORRADOR — Sujeto a revisión" cuando
    piar.estado == 'borrador'. Cuando 'aprobado', el encabezado muestra la
    fecha de aprobación en lugar de la marca.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor
    from documento import _set_cell_bg, _HEX_AZUL_OSC, _HEX_GRIS_CLR

    AZUL_OSCURO = RGBColor(0x1E, 0x40, 0xAF)
    AZUL_MEDIO = RGBColor(0x1D, 0x4E, 0xD8)
    AZUL_CLARO = RGBColor(0x93, 0xC5, 0xFD)
    BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
    GRIS = RGBColor(0x6B, 0x72, 0x80)
    ROJO_BORRADOR = RGBColor(0xC0, 0x39, 0x2B)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Franja superior
    tbl_top = doc.add_table(rows=1, cols=2)
    cl = tbl_top.cell(0, 0)
    _set_cell_bg(cl, _HEX_AZUL_OSC)
    pl = cl.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pl.paragraph_format.space_before = Pt(8)
    pl.paragraph_format.space_after = Pt(8)
    rl = pl.add_run(
        "  Asistente Pedagógico IA — PIAR"
    )
    rl.font.size = Pt(13); rl.font.bold = True; rl.font.color.rgb = BLANCO

    cr = tbl_top.cell(0, 1)
    _set_cell_bg(cr, _HEX_AZUL_OSC)
    pr = cr.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr.paragraph_format.space_before = Pt(8)
    pr.paragraph_format.space_after = Pt(8)
    rr = pr.add_run(f"Generado {datetime.now().strftime('%d/%m/%Y')}  ")
    rr.font.size = Pt(9); rr.font.color.rgb = AZUL_CLARO

    # Marca BORRADOR / APROBADO
    marca = doc.add_paragraph()
    marca.alignment = WD_ALIGN_PARAGRAPH.CENTER
    marca.paragraph_format.space_before = Pt(6)
    marca.paragraph_format.space_after = Pt(4)
    if piar.estado == "borrador":
        r_marca = marca.add_run("BORRADOR — Sujeto a revisión")
        r_marca.font.size = Pt(11); r_marca.font.bold = True
        r_marca.font.color.rgb = ROJO_BORRADOR
    else:
        fecha_aprob = piar.aprobado_en.strftime("%d/%m/%Y") if piar.aprobado_en else "—"
        r_marca = marca.add_run(f"APROBADO — {fecha_aprob}")
        r_marca.font.size = Pt(11); r_marca.font.bold = True
        r_marca.font.color.rgb = RGBColor(0x15, 0x80, 0x3D)

    # Metadatos
    inst = (getattr(docente, "institucion", None) or "Sin institución").strip()
    ciudad = (getattr(docente, "ciudad", None) or "").strip()
    inst_display = " — ".join(filter(None, [inst, ciudad]))

    tbl_meta = doc.add_table(rows=3, cols=2)
    meta = [
        [("Docente:", docente.nombre_completo), ("Institución:", inst_display)],
        [("Estudiante (código):", estudiante.codigo_estudiante),
         ("Grupo:", f"{grupo.nombre_grupo} · {grupo.grado} · {grupo.asignatura}")],
        [("Periodo:", f"{piar.periodo}  ·  Año {piar.anio}"),
         ("Versión:", f"v{piar.version}  ·  Estado: {piar.estado.upper()}")],
    ]
    for ri, fila in enumerate(meta):
        for ci, (label, valor) in enumerate(fila):
            cell = tbl_meta.cell(ri, ci)
            _set_cell_bg(cell, _HEX_GRIS_CLR)
            p = cell.paragraphs[0]
            p.paragraph_format.left_indent = Cm(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            r_lb = p.add_run(label + " ")
            r_lb.font.size = Pt(8); r_lb.font.bold = True; r_lb.font.color.rgb = AZUL_MEDIO
            r_val = p.add_run(valor)
            r_val.font.size = Pt(8); r_val.font.color.rgb = GRIS

    doc.add_paragraph()
    h = doc.add_heading("Plan Individual de Ajustes Razonables", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = AZUL_OSCURO
        run.font.size = Pt(16)

    marco = doc.add_paragraph()
    marco.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_marco = marco.add_run(
        "Decreto 1421 de 2017 · Ministerio de Educación Nacional de Colombia"
    )
    r_marco.font.size = Pt(9); r_marco.font.italic = True
    r_marco.font.color.rgb = GRIS

    # Secciones
    contenido = piar.contenido or {}
    for clave, titulo in _SECCION_LABEL.items():
        h_sec = doc.add_heading(titulo, level=2)
        for run in h_sec.runs:
            run.font.color.rgb = AZUL_OSCURO
            run.font.size = Pt(13)
        texto = contenido.get(clave) or _seccion_pendiente(clave)
        # Renderizado plano: preserva saltos de línea; sin parseo markdown
        # complejo (bullets, negritas) para MVP — se puede enriquecer luego.
        for parrafo in texto.split("\n\n"):
            if not parrafo.strip():
                continue
            p = doc.add_paragraph(parrafo.strip())
            p.paragraph_format.space_after = Pt(6)
            for r in p.runs:
                r.font.size = Pt(11)

    # Pie
    doc.add_paragraph()
    tbl_pie = doc.add_table(rows=1, cols=1)
    cp = tbl_pie.cell(0, 0)
    _set_cell_bg(cp, _HEX_GRIS_CLR)
    pp = cp.paragraphs[0]
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.paragraph_format.space_before = Pt(6)
    pp.paragraph_format.space_after = Pt(6)
    rp = pp.add_run(
        "  Documento generado con Asistente Pedagógico IA sobre base del Decreto 1421/2017. "
        "El template se ajustará al formato oficial del Instituto Manizales cuando esté disponible."
    )
    rp.font.size = Pt(8); rp.font.italic = True; rp.font.color.rgb = AZUL_MEDIO

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════
# ENDPOINTS
# ═══════════════════════════════════════════════════════════════

@router.post("/", response_model=PIAROut, status_code=201)
async def crear_piar(
    body: PIARCreateRequest,
    docente: Docente = Depends(get_current_docente),
    db: Session = Depends(get_db),
):
    """
    Crea un borrador de PIAR sintetizando la conversación PIAR existente
    entre el docente y el asistente para el estudiante indicado.

    Requiere que exista al menos un mensaje previo en modo PIAR para ese
    (grupo, estudiante). Si no hay conversación, devuelve 400.

    Consume 1 llamada al modelo (síntesis) — NO se contabiliza contra el
    rate limit diario aquí porque el rate limit se aplica en el socket
    handler para el chat, y esta síntesis es un turno diferenciado. Se
    documenta como decisión de MVP; en el próximo sprint puede integrarse
    al mismo contador si el owner lo pide.
    """
    # 1. Estudiante debe existir y pertenecer a un grupo del docente
    estudiante = db.query(Estudiante).filter(
        Estudiante.id_estudiante == body.id_estudiante,
    ).first()
    if not estudiante:
        raise HTTPException(status_code=404, detail="Estudiante no encontrado")
    grupo = db.query(Grupo).filter(
        Grupo.id_grupo == estudiante.id_grupo,
        Grupo.id_docente == docente.id_docente,
    ).first()
    if not grupo:
        raise HTTPException(status_code=404, detail="Estudiante no pertenece a un grupo tuyo")
    if not estudiante.tiene_piar:
        raise HTTPException(
            status_code=400,
            detail="El estudiante no tiene PIAR activo. Actívalo en la ficha del estudiante primero.",
        )

    # 2. Debe haber conversación previa en modo PIAR para este estudiante
    historial = (
        db.query(Mensaje)
        .filter(
            Mensaje.id_grupo == grupo.id_grupo,
            Mensaje.modo == MODO_PIAR,
            Mensaje.id_estudiante == estudiante.id_estudiante,
        )
        .order_by(Mensaje.timestamp.asc())
        .all()
    )
    if len(historial) == 0:
        raise HTTPException(
            status_code=400,
            detail=(
                "No hay conversación de PIAR aún para este estudiante. "
                "Iniciá el chat en modo PIAR antes de generar el borrador."
            ),
        )

    # 3. Sintetizar con Claude (o skeleton si falla)
    try:
        contenido = await _sintetizar_conversacion_a_json(
            docente, grupo, estudiante, historial,
        )
    except RuntimeError as exc:
        # Fallo del modelo — devolvemos skeleton + status distinto para debug
        raise HTTPException(status_code=502, detail=str(exc))

    # 4. Calcular versión y persistir
    anio = body.anio or grupo.anio_lectivo
    version = _next_version(db, estudiante.id_estudiante, grupo.id_grupo,
                            body.periodo, anio)
    piar = PIAR(
        id_estudiante=estudiante.id_estudiante,
        id_grupo=grupo.id_grupo,
        id_docente=docente.id_docente,
        periodo=body.periodo,
        anio=anio,
        version=version,
        contenido=contenido,
        estado="borrador",
    )
    db.add(piar)
    db.commit()
    db.refresh(piar)
    return piar


@router.get("/estudiante/{id_estudiante}", response_model=List[PIAROut])
def listar_por_estudiante(
    id_estudiante: str,
    docente: Docente = Depends(get_current_docente),
    db: Session = Depends(get_db),
):
    """
    Lista todos los PIARs (todas las versiones) del estudiante, ordenados
    del más reciente al más antiguo. Sólo del docente autenticado.
    """
    # Verificar que el estudiante es de un grupo del docente
    est = (
        db.query(Estudiante)
        .join(Grupo, Grupo.id_grupo == Estudiante.id_grupo)
        .filter(
            Estudiante.id_estudiante == id_estudiante,
            Grupo.id_docente == docente.id_docente,
        )
        .first()
    )
    if not est:
        raise HTTPException(status_code=404, detail="Estudiante no encontrado")

    piars = (
        db.query(PIAR)
        .filter(
            PIAR.id_estudiante == id_estudiante,
            PIAR.id_docente == docente.id_docente,
        )
        .order_by(PIAR.anio.desc(), PIAR.periodo.desc(), PIAR.version.desc())
        .all()
    )
    return piars


@router.put("/{piar_id}/aprobar", response_model=PIAROut)
def aprobar_piar(
    piar_id: str,
    docente: Docente = Depends(get_current_docente),
    db: Session = Depends(get_db),
):
    """
    Aprueba el PIAR: transición borrador → aprobado. Es inmutable después:
    reaprobar el mismo PIAR devuelve 409.

    Una vez aprobado, para hacer cambios el docente debe crear una nueva
    versión (POST /piar/ genera v+1 automáticamente).
    """
    piar = _piar_del_docente_o_404(piar_id, docente.id_docente, db)

    if piar.estado == "aprobado":
        raise HTTPException(
            status_code=409,
            detail=(
                "Este PIAR ya está aprobado. Para hacer cambios, generá una "
                "nueva versión desde el chat en modo PIAR."
            ),
        )

    piar.estado = "aprobado"
    piar.aprobado_en = datetime.utcnow()
    db.commit()
    db.refresh(piar)
    return piar


@router.get("/{piar_id}/docx")
def descargar_docx(
    piar_id: str,
    docente: Docente = Depends(get_current_docente),
    db: Session = Depends(get_db),
):
    """
    Genera el DOCX del PIAR on-demand y lo devuelve como attachment.
    Incluye marca 'BORRADOR — Sujeto a revisión' si el estado es borrador,
    o 'APROBADO — fecha' si ya fue aprobado.
    """
    piar = _piar_del_docente_o_404(piar_id, docente.id_docente, db)
    estudiante = db.query(Estudiante).filter(
        Estudiante.id_estudiante == piar.id_estudiante,
    ).first()
    grupo = db.query(Grupo).filter(Grupo.id_grupo == piar.id_grupo).first()
    if not estudiante or not grupo:
        raise HTTPException(status_code=500, detail="Estudiante o grupo del PIAR no encontrado (datos inconsistentes)")

    try:
        docx_bytes = _construir_piar_docx(docente, grupo, estudiante, piar)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Error generando DOCX del PIAR: {exc}")

    safe = re.sub(r"[^\w\s-]", "", estudiante.codigo_estudiante).strip().replace(" ", "_")[:40]
    filename = f"PIAR_{safe or 'estudiante'}_P{piar.periodo}_v{piar.version}.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )

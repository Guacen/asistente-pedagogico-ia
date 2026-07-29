"""
Template DOCX de Maestr.ia — layout de marca aplicado a contenido genérico.

Uso principal: `generar_piar_docx(secciones: dict, datos: dict) -> BytesIO`.

El diseño (encabezado + portada + cuerpo por secciones + pie + numeración)
está centralizado acá. El contenido de las secciones viene ya parseado
por sección (el modelo devolvió Markdown, `markdown_parser` lo dividió).

Paleta de marca (fuente de verdad: backend/branding.py):
  VERDE_OSCURO    #0B3D2E   títulos, wordmark "Maestr"
  VERDE_PRINCIPAL #1D9E75   subtítulos, líneas decorativas, ".ia"
  AMBAR           #F5B731   destacados
  VERDE_MENTA     #E1F5EE   fondo de header de tabla
  GRIS_PIE        #888888   footer neutro
"""
from __future__ import annotations

import io
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from markdown_parser import iter_inline_bold, parse_section_blocks


# ═══════════════════════════════════════════════════════════════
# PALETA — no importar de branding.py para no crear ciclos si
# branding.py crece; se replican los 5 hex que este template usa.
# ═══════════════════════════════════════════════════════════════

VERDE_OSCURO    = RGBColor(0x0B, 0x3D, 0x2E)
VERDE_PRINCIPAL = RGBColor(0x1D, 0x9E, 0x75)
AMBAR           = RGBColor(0xF5, 0xB7, 0x31)
VERDE_MENTA_HEX = "E1F5EE"   # sombreado de tabla (hex sin '#')
GRIS_PIE        = RGBColor(0x88, 0x88, 0x88)
BLANCO          = RGBColor(0xFF, 0xFF, 0xFF)
NEGRO_TEXTO     = RGBColor(0x1F, 0x25, 0x37)

FONT_FAMILY = "Calibri"   # Nunito no está garantizada en Word; Calibri sí.

TAGLINE = "Maestr.ia · Tu colega que conoce la ley"
LOGO_PATH_DEFAULT = Path(__file__).resolve().parent.parent / "frontend" / "assets" / "logo.png"


# ═══════════════════════════════════════════════════════════════
# HELPERS — celdas / bordes / párrafos / líneas decorativas
# ═══════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tbl_pr.append(borders)


def _add_horizontal_rule(paragraph, color: RGBColor, size_pt: float = 1.5) -> None:
    """
    Inserta una línea horizontal decorativa DENTRO de un párrafo usando
    el atributo `pBdr` (bottom). Grosor = size_pt * 8 (unidades del OOXML).
    """
    p_pr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(size_pt * 8)))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    pBdr.append(bottom)
    p_pr.append(pBdr)


def _run(paragraph, texto: str, *, bold: bool = False, size: int = 11,
         color: Optional[RGBColor] = None, italic: bool = False):
    r = paragraph.add_run(texto)
    r.font.name = FONT_FAMILY
    r.font.size = Pt(size)
    r.bold = bool(bold)
    r.italic = bool(italic)
    if color is not None:
        r.font.color.rgb = color
    return r


def _add_page_number_field(paragraph) -> None:
    """
    Inserta el campo dinámico { PAGE } que Word / LibreOffice actualizan
    con el número de página al abrir el documento.
    """
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)
    run.font.name = FONT_FAMILY
    run.font.size = Pt(8)
    run.font.color.rgb = GRIS_PIE


# ═══════════════════════════════════════════════════════════════
# HEADER + FOOTER — se aplican a todas las páginas
# ═══════════════════════════════════════════════════════════════

def _build_header(doc, logo_path: Optional[Path]) -> None:
    """
    Encabezado con: logo (si existe .png), tagline a la derecha, línea
    verde decorativa debajo. Si no hay logo, sólo el texto — no crash.
    """
    section = doc.sections[0]
    header = section.header

    tbl = header.add_table(rows=1, cols=2, width=Cm(17))
    tbl.autofit = False
    tbl.columns[0].width = Cm(6)
    tbl.columns[1].width = Cm(11)
    _remove_table_borders(tbl)

    # Izquierda: logo PNG (si existe) — python-docx no soporta SVG
    left = tbl.cell(0, 0)
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    left_p = left.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if logo_path and logo_path.exists():
        try:
            run = left_p.add_run()
            run.add_picture(str(logo_path), height=Cm(1.2))
        except Exception as exc:
            # PNG corrupto o formato no soportado — cae al texto.
            print(f"⚠️  add_picture falló ({exc}); usando fallback texto.")
            _run(left_p, "Maestr.ia", bold=True, size=13, color=VERDE_OSCURO)
    else:
        _run(left_p, "Maestr.ia", bold=True, size=13, color=VERDE_OSCURO)

    # Derecha: tagline
    right = tbl.cell(0, 1)
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right_p = right.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(right_p, TAGLINE, size=9, color=VERDE_PRINCIPAL, italic=True)

    # Línea decorativa horizontal verde principal debajo del header
    rule_p = header.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(2)
    rule_p.paragraph_format.space_after = Pt(0)
    _add_horizontal_rule(rule_p, VERDE_PRINCIPAL, size_pt=1.5)


def _build_footer(doc) -> None:
    """
    Pie con "Maestr.ia · Generado con IA · <fecha>" a la izquierda y
    número de página a la derecha. Aplica a todas las páginas.
    """
    section = doc.sections[0]
    footer = section.footer

    tbl = footer.add_table(rows=1, cols=2, width=Cm(17))
    tbl.autofit = False
    tbl.columns[0].width = Cm(14)
    tbl.columns[1].width = Cm(3)
    _remove_table_borders(tbl)

    fecha_hoy = datetime.utcnow().strftime("%d/%m/%Y")
    left = tbl.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(left, f"Maestr.ia · Generado con IA · {fecha_hoy}",
         size=8, color=GRIS_PIE, italic=True)

    right = tbl.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(right, "Página ", size=8, color=GRIS_PIE)
    _add_page_number_field(right)


# ═══════════════════════════════════════════════════════════════
# PORTADA — título + subtítulo + tabla de datos
# ═══════════════════════════════════════════════════════════════

def _build_portada(
    doc,
    titulo: str,
    subtitulo: str,
    datos: dict[str, Any],
) -> None:
    doc.add_paragraph()   # espacio antes del título

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_titulo, titulo, bold=True, size=18, color=VERDE_OSCURO)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_sub, subtitulo, size=11, color=VERDE_PRINCIPAL, italic=True)

    doc.add_paragraph()   # espacio antes de la tabla de datos

    # Tabla 2 columnas: campo | valor
    campos = [
        ("Nombre", datos.get("nombre", "—")),
        ("Grado",  datos.get("grado",  "—")),
        ("Docente", datos.get("docente", "—")),
        ("Fecha",  datos.get("fecha", datetime.utcnow().strftime("%d/%m/%Y"))),
    ]
    tbl = doc.add_table(rows=len(campos), cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Cm(4)
    tbl.columns[1].width = Cm(13)
    for i, (k, v) in enumerate(campos):
        c_label = tbl.cell(i, 0)
        _set_cell_bg(c_label, VERDE_MENTA_HEX)
        p_label = c_label.paragraphs[0]
        p_label.paragraph_format.left_indent = Cm(0.2)
        _run(p_label, k, bold=True, size=10, color=VERDE_OSCURO)

        c_val = tbl.cell(i, 1)
        p_val = c_val.paragraphs[0]
        p_val.paragraph_format.left_indent = Cm(0.2)
        _run(p_val, str(v), size=10, color=NEGRO_TEXTO)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# CUERPO — una sección por heading
# ═══════════════════════════════════════════════════════════════

def _render_paragraph_inline(doc_paragraph, texto: str) -> None:
    """Emite runs con negrita para **texto** en línea."""
    for segmento, is_bold in iter_inline_bold(texto):
        _run(doc_paragraph, segmento, bold=is_bold, size=11, color=NEGRO_TEXTO)


def _add_seccion(doc, titulo: str, contenido_md: str) -> None:
    # Heading de sección (H2 stylizado con color de marca)
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(2)
    _run(h, titulo, bold=True, size=13, color=VERDE_OSCURO)

    # Línea decorativa fina bajo el heading
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after = Pt(6)
    _add_horizontal_rule(rule, VERDE_PRINCIPAL, size_pt=0.75)

    if not contenido_md.strip():
        p_vacio = doc.add_paragraph()
        _run(p_vacio, "[PENDIENTE — sin información]",
             italic=True, size=10, color=GRIS_PIE)
        return

    for tipo, payload in parse_section_blocks(contenido_md):
        if tipo == "paragraph":
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            _render_paragraph_inline(p, payload)
        elif tipo == "bullet":
            # ListBullet: estilo builtin de Word "List Bullet"
            try:
                p = doc.add_paragraph(style="List Bullet")
            except KeyError:
                # Templates minimalistas de python-docx a veces no tienen
                # el estilo List Bullet. Fallback a párrafo con "• ".
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                _run(p, "• ", size=11, color=VERDE_PRINCIPAL)
            _render_paragraph_inline(p, payload)
        elif tipo == "heading3":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            _run(p, payload, bold=True, size=11, color=VERDE_PRINCIPAL)
        elif tipo == "table":
            _add_table_from_data(doc, payload)


def _add_table_from_data(doc, filas: list[list[str]]) -> None:
    if not filas:
        return
    ncols = max(len(f) for f in filas)
    tbl = doc.add_table(rows=len(filas), cols=ncols)
    tbl.style = "Table Grid"
    for i, fila in enumerate(filas):
        for j in range(ncols):
            celda = tbl.cell(i, j)
            texto = fila[j] if j < len(fila) else ""
            p = celda.paragraphs[0]
            p.paragraph_format.left_indent = Cm(0.15)
            if i == 0:
                _set_cell_bg(celda, VERDE_MENTA_HEX)
                _run(p, texto, bold=True, size=10, color=VERDE_OSCURO)
            else:
                _run(p, texto, size=10, color=NEGRO_TEXTO)
    doc.add_paragraph()   # aire después de la tabla


# ═══════════════════════════════════════════════════════════════
# API PÚBLICA
# ═══════════════════════════════════════════════════════════════

def generar_piar_docx(
    secciones: dict[str, str],
    datos_estudiante: dict[str, Any],
    *,
    logo_path: Optional[Path] = None,
) -> io.BytesIO:
    """
    Renderiza el DOCX del PIAR desde `secciones` (dict título→markdown)
    y los datos del estudiante para la portada.

    Args:
        secciones: dict con {"Datos del estudiante": "...", "Barreras...": "...", ...}
                   El orden del dict determina el orden en el documento.
        datos_estudiante: {"nombre", "grado", "docente", "fecha"}. Todos
                          strings; los faltantes se rellenan con "—".
        logo_path: ruta al PNG del logo (Path). Si es None o no existe,
                   el header usa fallback texto "Maestr.ia".

    Retorna un `io.BytesIO` con el documento listo para stream/send.
    """
    doc = Document()

    # Márgenes conservadores estilo A4 pedagógico
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        # Espacio para el header + separator
        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.8)

    _build_header(doc, logo_path if logo_path is not None else LOGO_PATH_DEFAULT)
    _build_footer(doc)

    _build_portada(
        doc,
        titulo="Plan Individual de Ajustes Razonables (PIAR)",
        subtitulo="Decreto 1421 de 2017 — Educación Inclusiva",
        datos=datos_estudiante,
    )

    for titulo, contenido in secciones.items():
        _add_seccion(doc, titulo, contenido or "")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

"""
Generación de documentos DOCX con encabezado institucional.

Endpoint: POST /api/grupos/{grupo_id}/generar-documento
Body   : { "contenido_md": "...", "titulo": "..." }
Returns: archivo .docx descargable

Flujo:
  1. El frontend acumula el Markdown de la respuesta IA en el chat.
  2. Detecta que es un documento (contiene encabezados/listas).
  3. POST a este endpoint con el MD y el título.
  4. El backend convierte MD → DOCX con encabezado institucional.
  5. El browser descarga el archivo.
"""

import io
import re
from datetime import datetime
from typing import List, Optional

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from auth import verify_trial_active
from database import get_db
from models import Calificacion, Docente, Estudiante, EvaluacionColumna, Grupo

router = APIRouter(prefix="/api", tags=["documentos"])


# ════════════════════════════════════════════════════════════════
# SCHEMA
# ════════════════════════════════════════════════════════════════

class GenerarDocumentoRequest(BaseModel):
    contenido_md: str
    titulo: str = "Documento"


# ════════════════════════════════════════════════════════════════
# HELPERS DOCX — encabezado institucional estilo Teachy
# ════════════════════════════════════════════════════════════════

def _apply_brand_heading_styles(doc) -> None:
    """
    Redefine los estilos globales Heading 1/2/3 con la paleta Maestr.ia.
    Aplica al render actual Y a cualquier heading que el docente agregue
    después editando el DOCX en Word — coherencia total.

    Heading 1: 16pt bold #0B3D2E
    Heading 2: 13pt bold #0B3D2E
    Heading 3: 11pt bold #1D9E75

    Bug 3 del ticket "docx-5-bugs": antes Heading 3 heredaba #4F81BD
    (azul default de python-docx) porque este generador nunca redefinía
    los estilos — solo forzaba color a nivel de run.
    """
    from docx.shared import Pt, RGBColor
    VERDE_OSCURO = RGBColor(0x0B, 0x3D, 0x2E)
    VERDE_PRIM   = RGBColor(0x1D, 0x9E, 0x75)
    specs = [
        ('Heading 1', 16, VERDE_OSCURO),
        ('Heading 2', 13, VERDE_OSCURO),
        ('Heading 3', 11, VERDE_PRIM),
    ]
    for name, size, color in specs:
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = 'Calibri'
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color


def _docx_bytes(md: str, titulo: str, docente: Docente, grupo: Grupo) -> bytes:
    """
    Construye el DOCX completo:
      - Encabezado institucional con marca de la plataforma
      - Cuerpo desde Markdown
      - Pie de página con branding
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    # ── Paleta ──────────────────────────────────────────────────
    AZUL_OSCURO = RGBColor(0x0B, 0x3D, 0x2E)   # azul-800 Tailwind
    AZUL_MEDIO  = RGBColor(0x1D, 0x9E, 0x75)   # azul-700
    AZUL_CLARO  = RGBColor(0xB6, 0xE5, 0xD1)   # azul-300
    BLANCO      = RGBColor(0xFF, 0xFF, 0xFF)
    GRIS        = RGBColor(0x6B, 0x72, 0x80)

    HEX_AZUL_OSC = '0B3D2E'
    HEX_AZUL_CLR = 'E1F5EE'
    HEX_GRIS_CLR = 'F3F4F6'

    # ── Helpers ─────────────────────────────────────────────────
    def set_cell_bg(cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def set_cell_border_none(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            bd = OxmlElement(f'w:{side}')
            bd.set(qn('w:val'), 'nil')
            tcBorders.append(bd)
        tcPr.append(tcBorders)

    def no_border_table(tbl):
        tbl.style = 'Table Grid'
        for row in tbl.rows:
            for cell in row.cells:
                set_cell_border_none(cell)

    # ── Documento ────────────────────────────────────────────────
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── section.header + section.footer con marca (todas las páginas) ──
    # La "franja superior" que sigue abajo vive en el BODY del documento
    # (aparece sólo en la primera página); el header/footer de Word se
    # aplican a TODAS las páginas y son lo que python-docx detecta al
    # inspeccionar doc.sections[0].header/.footer. Antes quedaban vacíos.
    # Path del logo — desde backend/documento.py: parent (=backend) +
    # frontend/assets/logo.png. Calculado con Path.__file__.parent para
    # que funcione desde cualquier cwd (Railway / dev / tests).
    from pathlib import Path
    _LOGO_PATH = Path(__file__).resolve().parent / 'frontend' / 'assets' / 'logo.png'

    for section in doc.sections:
        # Header: logo a la izquierda + tagline a la derecha.
        # Tabla 1x2 sin bordes — así python-docx detecta el drawing
        # dentro del section.header y el layout queda pareado.
        header = section.header
        # Limpiar el paragraph vacío inicial que python-docx crea, sino
        # queda un espacio en blanco arriba del contenido real del header.
        if header.paragraphs and not header.paragraphs[0].text.strip():
            # No lo removemos (rompe el elemento); solo lo usamos como anchor
            # invisible cerca del top. Toda la marca va en la tabla siguiente.
            pass
        h_tbl = header.add_table(rows=1, cols=2, width=Cm(17))
        h_tbl.autofit = False
        h_tbl.columns[0].width = Cm(4)
        h_tbl.columns[1].width = Cm(13)

        # Izquierda: logo PNG si existe, sino fallback texto
        left_cell = h_tbl.cell(0, 0)
        left_p = left_cell.paragraphs[0]
        left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if _LOGO_PATH.exists():
            try:
                left_p.add_run().add_picture(str(_LOGO_PATH), height=Cm(1.0))
            except Exception as e:
                print(f"[documento.py] logo header fail: {e} — path={_LOGO_PATH}")
                _fb = left_p.add_run('Maestr.ia')
                _fb.font.name = 'Calibri'; _fb.font.bold = True
                _fb.font.size = Pt(11); _fb.font.color.rgb = AZUL_OSCURO
        else:
            print(f"[documento.py] logo no encontrado: {_LOGO_PATH}")
            _fb = left_p.add_run('Maestr.ia')
            _fb.font.name = 'Calibri'; _fb.font.bold = True
            _fb.font.size = Pt(11); _fb.font.color.rgb = AZUL_OSCURO

        # Derecha: tagline
        right_cell = h_tbl.cell(0, 1)
        right_p = right_cell.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        h_run = right_p.add_run('Maestr.ia · Tu colega que conoce la ley')
        h_run.font.name = 'Calibri'
        h_run.font.size = Pt(8)
        h_run.font.italic = True
        h_run.font.color.rgb = AZUL_MEDIO   # verde primario post-rebrand

        # Footer: marca + fecha, centrado
        f_p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        f_run = f_p.add_run(
            f'Maestr.ia · Tu colega que conoce la ley · Generado con IA · '
            f'{datetime.now().strftime("%d/%m/%Y")}'
        )
        f_run.font.name = 'Calibri'
        f_run.font.size = Pt(8)
        f_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ══════════════════════════════════════════════════════════════
    # FRANJA SUPERIOR — verde sólido con marca
    # ══════════════════════════════════════════════════════════════
    tbl_top = doc.add_table(rows=1, cols=2)
    no_border_table(tbl_top)
    tbl_top.columns[0].width = Cm(10)
    tbl_top.columns[1].width = Cm(8)

    # Celda izquierda: nombre de la plataforma
    cl = tbl_top.cell(0, 0)
    set_cell_bg(cl, HEX_AZUL_OSC)
    pl = cl.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pl.paragraph_format.left_indent  = Cm(0.4)
    pl.paragraph_format.space_before = Pt(8)
    pl.paragraph_format.space_after  = Pt(8)
    rl = pl.add_run('Maestr.ia')
    rl.font.size  = Pt(14)
    rl.font.bold  = True
    rl.font.color.rgb = BLANCO

    # Celda derecha: fecha
    cr = tbl_top.cell(0, 1)
    set_cell_bg(cr, HEX_AZUL_OSC)
    pr = cr.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr.paragraph_format.right_indent = Cm(0.4)
    pr.paragraph_format.space_before = Pt(8)
    pr.paragraph_format.space_after  = Pt(8)
    rr = pr.add_run(f'Generado: {datetime.now().strftime("%d/%m/%Y  %H:%M")}')
    rr.font.size      = Pt(9)
    rr.font.color.rgb = AZUL_CLARO

    # ══════════════════════════════════════════════════════════════
    # TABLA DE METADATOS (2 filas × 2 cols)
    # ══════════════════════════════════════════════════════════════
    tbl_meta = doc.add_table(rows=2, cols=2)
    no_border_table(tbl_meta)

    inst = (getattr(docente, 'institucion', None) or '').strip()
    ciudad = (getattr(docente, 'ciudad', None) or '').strip()
    inst_display = ' — '.join(filter(None, [inst or 'Sin institución', ciudad]))

    meta = [
        [
            ('Docente:', docente.nombre_completo),
            ('Generado por:', 'Maestr.ia'),
        ],
        [
            ('Institución:', inst_display),
            ('Grupo:', f'{grupo.nombre_grupo} · {grupo.grado} · {grupo.asignatura} · P{grupo.periodo_actual}'),
        ],
    ]

    for ri, fila in enumerate(meta):
        for ci, (label, valor) in enumerate(fila):
            cell = tbl_meta.cell(ri, ci)
            set_cell_bg(cell, HEX_GRIS_CLR)
            p = cell.paragraphs[0]
            p.paragraph_format.left_indent  = Cm(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            r_label = p.add_run(label + ' ')
            r_label.font.size      = Pt(8)
            r_label.font.bold      = True
            r_label.font.color.rgb = AZUL_MEDIO
            r_val = p.add_run(valor)
            r_val.font.size      = Pt(8)
            r_val.font.color.rgb = GRIS

    # Espacio
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)

    # ══════════════════════════════════════════════════════════════
    # TÍTULO DEL DOCUMENTO
    # ══════════════════════════════════════════════════════════════
    h_titulo = doc.add_heading(titulo, level=1)
    h_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h_titulo.runs:
        run.font.color.rgb = AZUL_OSCURO
        run.font.size = Pt(18)

    # Espaciador entre título y cuerpo (antes había una fila de 85 ─,
    # bug 2 del ticket docx-5-bugs). El space_after del párrafo del
    # título ya provee la respiración; agregamos solo un párrafo vacío
    # de altura mínima para no perder el aire.
    _sp = doc.add_paragraph()
    _sp.paragraph_format.space_before = Pt(0)
    _sp.paragraph_format.space_after  = Pt(8)

    # ══════════════════════════════════════════════════════════════
    # CUERPO — Markdown → DOCX
    # ══════════════════════════════════════════════════════════════
    # Aplicar estilos de marca a Heading 1/2/3 globales (bug 3): así
    # todo add_heading dentro de _render_md hereda los colores verdes
    # aunque no forcemos color en run explícito.
    _apply_brand_heading_styles(doc)
    # Skip primer H1 del markdown si duplica el título de la portada.
    _render_md(doc, md, AZUL_OSCURO, skip_h1_titulo=titulo)

    # ══════════════════════════════════════════════════════════════
    # PIE — franja de marca
    # ══════════════════════════════════════════════════════════════
    pie_sp = doc.add_paragraph()
    pie_sp.paragraph_format.space_before = Pt(16)
    pie_sp.paragraph_format.space_after  = Pt(0)

    tbl_pie = doc.add_table(rows=1, cols=1)
    no_border_table(tbl_pie)
    cp = tbl_pie.cell(0, 0)
    set_cell_bg(cp, HEX_AZUL_CLR)
    pp = cp.paragraphs[0]
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.paragraph_format.space_before = Pt(6)
    pp.paragraph_format.space_after  = Pt(6)
    rp = pp.add_run(
        f'Generado con Maestr.ia  •  '
        f'{datetime.now().strftime("%d/%m/%Y")}  •  '
        'Documento de uso educativo — No editar el encabezado'
    )
    rp.font.size      = Pt(7)
    rp.font.italic    = True
    rp.font.color.rgb = AZUL_MEDIO

    # ── Serializar ───────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


_SEP_CHARS = set('─-—=_')


def _es_linea_separador(s: str) -> bool:
    """
    Detecta líneas que son solo caracteres decorativos de separación
    (guiones ASCII o box-drawing U+2500). Bug 2: antes se renderizaban
    como texto o como una línea gris de 80 ─; el ticket pide eliminarlas
    y confiar en el space_after de los headings.
    """
    st = s.strip()
    return len(st) > 5 and set(st) <= _SEP_CHARS


def _render_md(doc, md: str, h1_color, *, skip_h1_titulo: str = ''):
    """
    Parsea Markdown básico y lo escribe en el documento con python-docx.
    Soporta: # encabezados, **negrita**, *itálica*, listas, código,
    tablas Markdown (bug 5), párrafos. Salta el primer H1 si su texto
    coincide con `skip_h1_titulo` (bug 4 — evita título duplicado con
    la portada).
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor
    from markdown_parser import _parse_fila_tabla, _es_separador_tabla, _TABLE_ROW_RE

    AZUL = h1_color              # nombre legacy — ahora es VERDE_OSCURO
    VERDE_OSCURO = h1_color
    VERDE_PRIM = RGBColor(0x1D, 0x9E, 0x75)
    VERDE_MENTA_HEX = 'E1F5EE'
    skip_titulo_norm = (skip_h1_titulo or '').strip().lower()
    ya_saltamos_h1 = False

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        s = raw.strip()

        # ── Bloque de código ````lang ... ``` ────────────────────
        if s.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent  = Pt(20)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(4)
                r = p.add_run('\n'.join(code_lines))
                r.font.name = 'Courier New'
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            i += 1  # saltar cierre ```
            continue

        # ── Encabezados ──────────────────────────────────────────
        m = re.match(r'^(#{1,6})\s+(.+)', s)
        if m:
            level  = min(len(m.group(1)), 4)
            text   = _strip_inline(m.group(2))
            # Bug 4: skip primer H1 si coincide con el título de portada
            if (level == 1 and not ya_saltamos_h1
                    and skip_titulo_norm
                    and text.strip().lower() == skip_titulo_norm):
                ya_saltamos_h1 = True
                i += 1
                continue
            h = doc.add_heading(text, level=level)
            # Forzar colores a nivel de RUN (bug 3): el estilo global
            # ya está redefinido en verde, pero add_heading a veces
            # inserta runs con color heredado del template default.
            if level == 1:
                for r in h.runs:
                    r.font.color.rgb = VERDE_OSCURO
                    r.font.size = Pt(16)
            elif level == 2:
                for r in h.runs:
                    r.font.color.rgb = VERDE_OSCURO
                    r.font.size = Pt(13)
            elif level == 3:
                for r in h.runs:
                    r.font.color.rgb = VERDE_PRIM
                    r.font.size = Pt(11)
            elif level == 4:
                for r in h.runs:
                    r.font.color.rgb = VERDE_PRIM
                    r.font.size = Pt(10)
            i += 1
            continue

        # ── Tabla Markdown (bug 5) ───────────────────────────────
        # Si la línea empieza con `|` y termina con `|`, y la siguiente
        # es un separador |---|---|, procesamos el bloque entero como
        # tabla real de Word.
        if _TABLE_ROW_RE.match(s) and i + 1 < len(lines) and _es_separador_tabla(lines[i + 1]):
            filas = []
            header = _parse_fila_tabla(s)
            if header:
                filas.append(header)
            i += 2  # saltar header + separador
            while i < len(lines) and _TABLE_ROW_RE.match(lines[i].strip()):
                fila = _parse_fila_tabla(lines[i].strip())
                if fila:
                    filas.append(fila)
                i += 1
            if filas:
                ncols = max(len(f) for f in filas)
                tbl = doc.add_table(rows=len(filas), cols=ncols)
                tbl.style = 'Table Grid'
                for ri, fila in enumerate(filas):
                    for ci in range(ncols):
                        celda = tbl.cell(ri, ci)
                        texto = fila[ci] if ci < len(fila) else ''
                        p_c = celda.paragraphs[0]
                        p_c.paragraph_format.left_indent = Cm(0.15)
                        if ri == 0:
                            _set_cell_bg(celda, VERDE_MENTA_HEX)
                            r_c = p_c.add_run(texto)
                            r_c.font.name = 'Calibri'; r_c.font.bold = True
                            r_c.font.size = Pt(10); r_c.font.color.rgb = VERDE_OSCURO
                        else:
                            r_c = p_c.add_run(texto)
                            r_c.font.name = 'Calibri'; r_c.font.size = Pt(10)
                doc.add_paragraph()
            continue

        # ── Lista viñetas ────────────────────────────────────────
        m_bul = re.match(r'^[-*+]\s+(.+)', s)
        if m_bul:
            p = doc.add_paragraph(style='List Bullet')
            _inline_runs(p, m_bul.group(1))
            i += 1
            continue

        # ── Lista numerada ───────────────────────────────────────
        m_num = re.match(r'^\d+[.)]\s+(.+)', s)
        if m_num:
            p = doc.add_paragraph(style='List Number')
            _inline_runs(p, m_num.group(1))
            i += 1
            continue

        # ── Separadores decorativos (bug 2) ──────────────────────
        # Antes: cualquier línea de --- se renderizaba como 80 ─ grises.
        # Ahora: se ignoran completamente — el space_after de headings
        # ya provee la respiración visual entre secciones.
        if _es_linea_separador(s):
            i += 1
            continue
        if re.match(r'^[-*_]{3,}$', s):
            i += 1
            continue

        # ── Línea vacía → espacio ────────────────────────────────
        if not s:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            i += 1
            continue

        # ── Párrafo normal ───────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _inline_runs(p, s)
        i += 1


def _strip_inline(text: str) -> str:
    """Quita marcadores MD inline para texto plano."""
    t = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    t = re.sub(r'__(.+?)__', r'\1', t)
    t = re.sub(r'\*(.+?)\*', r'\1', t)
    t = re.sub(r'_(.+?)_', r'\1', t)
    t = re.sub(r'`(.+?)`', r'\1', t)
    t = re.sub(r'~~(.+?)~~', r'\1', t)
    return t


def _inline_runs(paragraph, text: str):
    """
    Agrega runs al párrafo respetando **negrita** e *itálica* inline.
    Tokeniza el texto y produce runs separados con el formato correcto.
    """
    from docx.shared import Pt

    # Tokenizar: separar en segmentos de texto con formato
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|__.*?__|_.*?_|`.*?`)', text)
    for token in tokens:
        if not token:
            continue
        m_bold   = re.match(r'\*\*(.+?)\*\*', token) or re.match(r'__(.+?)__', token)
        m_italic = re.match(r'\*(.+?)\*', token) or re.match(r'_(.+?)_', token)
        m_code   = re.match(r'`(.+?)`', token)
        if m_bold:
            r = paragraph.add_run(m_bold.group(1))
            r.bold = True
        elif m_italic:
            r = paragraph.add_run(m_italic.group(1))
            r.italic = True
        elif m_code:
            r = paragraph.add_run(m_code.group(1))
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
        else:
            paragraph.add_run(token)


# ════════════════════════════════════════════════════════════════
# ENDPOINT
# ════════════════════════════════════════════════════════════════

@router.post("/grupos/{grupo_id}/generar-documento")
async def generar_documento(
    grupo_id: str,
    body: GenerarDocumentoRequest,
    docente: Docente = Depends(verify_trial_active),
    db: Session = Depends(get_db),
):
    """
    Recibe Markdown de la IA, lo convierte a DOCX con encabezado institucional
    y devuelve el archivo para descarga directa.
    """
    grupo = db.query(Grupo).filter(
        Grupo.id_grupo == grupo_id,
        Grupo.id_docente == docente.id_docente,
    ).first()
    if not grupo:
        raise HTTPException(status_code=404, detail="Grupo no encontrado")

    titulo = (body.titulo or 'Documento').strip()[:120]

    try:
        docx_bytes = _docx_bytes(
            md=body.contenido_md,
            titulo=titulo,
            docente=docente,
            grupo=grupo,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generando DOCX: {e}")

    # Nombre de archivo seguro
    safe = re.sub(r'[^\w\s-]', '', titulo).strip().replace(' ', '_')[:60] or 'documento'
    filename = f"{safe}.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )


# ════════════════════════════════════════════════════════════════
# BOLETINES DE CALIFICACIONES
# ════════════════════════════════════════════════════════════════
#
# Endpoints:
#   GET /api/grupos/{grupo_id}/boletin/{estudiante_id}?periodo=N
#     → boletín individual DOCX
#   GET /api/grupos/{grupo_id}/boletin?periodo=N
#     → boletín consolidado del grupo (todos los estudiantes)
#
# Escala colombiana (misma que grupos.html tab Calificaciones):
#   ≥ 3.5 → Aprobado (verde)
#   3.0 – 3.4 → En riesgo (amarillo)
#   < 3.0 → Reprobado (rojo)
#
# Los helpers de estilo se duplican al top-level para no arriesgar el
# generador de documentos IA existente (_docx_bytes queda intacto).

_UMBRAL_APROBADO = 3.5
_UMBRAL_APROBACION_MIN = 3.0

# Paleta (mismos hex que _docx_bytes)
_HEX_AZUL_OSC = '0B3D2E'
_HEX_AZUL_CLR = 'E1F5EE'
_HEX_GRIS_CLR = 'F3F4F6'
_HEX_VERDE    = 'DCFCE7'
_HEX_AMBAR    = 'FEF9C3'
_HEX_ROJO     = 'FEE2E2'


def _promedio_ponderado(notas_por_columna: list[tuple[Optional[float], Optional[float]]]) -> Optional[float]:
    """
    notas_por_columna = [(valor, porcentaje_de_columna), ...]
    Peso = porcentaje de la columna si está definido y >0; si no, peso = 1.
    Devuelve None si no hay notas con valor.
    """
    acc = 0.0
    peso = 0.0
    for valor, porcentaje in notas_por_columna:
        if valor is None:
            continue
        w = porcentaje if (isinstance(porcentaje, (int, float)) and porcentaje > 0) else 1.0
        acc += float(valor) * w
        peso += w
    if peso <= 0:
        return None
    return round(acc / peso * 10) / 10


def _estado_desde_promedio(prom: Optional[float]) -> tuple[str, str]:
    """Devuelve (etiqueta, hex_bg) según la escala del docente."""
    if prom is None:
        return ('Sin notas', _HEX_GRIS_CLR)
    if prom >= _UMBRAL_APROBADO:
        return ('Aprobado', _HEX_VERDE)
    if prom >= _UMBRAL_APROBACION_MIN:
        return ('En riesgo', _HEX_AMBAR)
    return ('Reprobado', _HEX_ROJO)


def _set_cell_bg(cell, hex_color: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _boletin_encabezado(doc, docente: Docente, grupo: Grupo, subtitulo: str, periodo: int):
    """Encabezado institucional del boletín — mismo look que documentos IA."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    AZUL_OSCURO = RGBColor(0x0B, 0x3D, 0x2E)
    AZUL_MEDIO  = RGBColor(0x1D, 0x9E, 0x75)
    AZUL_CLARO  = RGBColor(0xB6, 0xE5, 0xD1)
    BLANCO      = RGBColor(0xFF, 0xFF, 0xFF)
    GRIS        = RGBColor(0x6B, 0x72, 0x80)

    # Franja superior
    tbl_top = doc.add_table(rows=1, cols=2)
    cl = tbl_top.cell(0, 0)
    _set_cell_bg(cl, _HEX_AZUL_OSC)
    pl = cl.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pl.paragraph_format.space_before = Pt(8)
    pl.paragraph_format.space_after  = Pt(8)
    rl = pl.add_run('  Maestr.ia · Boletín de Calificaciones')
    rl.font.size = Pt(13); rl.font.bold = True; rl.font.color.rgb = BLANCO

    cr = tbl_top.cell(0, 1)
    _set_cell_bg(cr, _HEX_AZUL_OSC)
    pr = cr.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr.paragraph_format.space_before = Pt(8)
    pr.paragraph_format.space_after  = Pt(8)
    rr = pr.add_run(f'Generado {datetime.now().strftime("%d/%m/%Y")}  ')
    rr.font.size = Pt(9); rr.font.color.rgb = AZUL_CLARO

    # Metadatos
    inst = (getattr(docente, 'institucion', None) or 'Sin institución').strip()
    ciudad = (getattr(docente, 'ciudad', None) or '').strip()
    inst_display = ' — '.join(filter(None, [inst, ciudad]))

    tbl_meta = doc.add_table(rows=2, cols=2)
    meta = [
        [('Docente:', docente.nombre_completo), ('Institución:', inst_display)],
        [('Grupo:', f'{grupo.nombre_grupo} · {grupo.grado} · {grupo.asignatura}'),
         ('Periodo:', f'{periodo}  ·  Año {grupo.anio_lectivo}')],
    ]
    for ri, fila in enumerate(meta):
        for ci, (label, valor) in enumerate(fila):
            cell = tbl_meta.cell(ri, ci)
            _set_cell_bg(cell, _HEX_GRIS_CLR)
            p = cell.paragraphs[0]
            p.paragraph_format.left_indent  = Cm(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            r_lb = p.add_run(label + ' ')
            r_lb.font.size = Pt(8); r_lb.font.bold = True; r_lb.font.color.rgb = AZUL_MEDIO
            r_val = p.add_run(valor)
            r_val.font.size = Pt(8); r_val.font.color.rgb = GRIS

    # Subtítulo
    doc.add_paragraph()
    h = doc.add_heading(subtitulo, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = AZUL_OSCURO
        run.font.size = Pt(16)


def _boletin_tabla_estudiante(doc, columnas, notas_por_col: dict, promedio: Optional[float]):
    """
    Escribe la tabla de calificaciones de un estudiante:
    fila 1: encabezados (evaluación, tipo, %, nota)
    filas 2..n: datos de cada columna
    fila final: Definitiva con color por estado
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    AZUL_OSCURO = RGBColor(0x0B, 0x3D, 0x2E)
    NEGRO = RGBColor(0x11, 0x18, 0x27)

    tbl = doc.add_table(rows=1 + len(columnas) + 1, cols=4)
    tbl.style = 'Table Grid'

    # Encabezados
    headers = ['Evaluación', 'Tipo', '% Peso', 'Nota']
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        _set_cell_bg(c, _HEX_AZUL_OSC)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Datos por columna
    for idx, col in enumerate(columnas, start=1):
        row = tbl.rows[idx]
        row.cells[0].text = col.nombre or '—'
        row.cells[1].text = (col.tipo or '—').capitalize()
        row.cells[2].text = f'{col.porcentaje:.0f}%' if col.porcentaje else '—'
        nota = notas_por_col.get(col.id_columna)
        row.cells[3].text = f'{nota:.1f}' if nota is not None else '—'
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9); r.font.color.rgb = NEGRO
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Definitiva
    row_def = tbl.rows[-1]
    _set_cell_bg(row_def.cells[0], _HEX_GRIS_CLR)
    _set_cell_bg(row_def.cells[1], _HEX_GRIS_CLR)
    _set_cell_bg(row_def.cells[2], _HEX_GRIS_CLR)
    etiqueta, hex_estado = _estado_desde_promedio(promedio)
    _set_cell_bg(row_def.cells[3], hex_estado)

    row_def.cells[0].text = ''
    p_def = row_def.cells[0].paragraphs[0]
    r_def = p_def.add_run('DEFINITIVA')
    r_def.font.bold = True; r_def.font.size = Pt(10); r_def.font.color.rgb = AZUL_OSCURO

    row_def.cells[1].text = etiqueta
    for r in row_def.cells[1].paragraphs[0].runs:
        r.font.bold = True; r.font.size = Pt(9)
    row_def.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row_def.cells[2].text = '100%' if columnas else '—'
    row_def.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row_def.cells[3].text = f'{promedio:.1f}' if promedio is not None else '—'
    for r in row_def.cells[3].paragraphs[0].runs:
        r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = AZUL_OSCURO
    row_def.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _boletin_pie_leyenda(doc):
    """Pie del boletín con leyenda de escala + branding."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    AZUL_MEDIO = RGBColor(0x1D, 0x9E, 0x75)

    doc.add_paragraph()
    tbl_leg = doc.add_table(rows=1, cols=1)
    cell = tbl_leg.cell(0, 0)
    _set_cell_bg(cell, _HEX_GRIS_CLR)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(
        '  Escala:  ≥ 3.5 Aprobado   ·   3.0 – 3.4 En riesgo   ·   < 3.0 Reprobado    '
        '·   Promedio ponderado por peso de columna (1.0 si no tiene peso definido).'
    )
    r.font.size = Pt(8); r.font.italic = True; r.font.color.rgb = AZUL_MEDIO


def _cargar_datos_boletin(db: Session, grupo: Grupo, periodo: int):
    """
    Devuelve (columnas_del_periodo, {est_id: {col_id: valor}}).
    """
    columnas = (
        db.query(EvaluacionColumna)
        .filter(
            EvaluacionColumna.id_grupo == grupo.id_grupo,
            EvaluacionColumna.periodo == periodo,
        )
        .order_by(EvaluacionColumna.orden, EvaluacionColumna.nombre)
        .all()
    )
    col_ids = {c.id_columna for c in columnas}
    notas_por_est: dict[str, dict[str, float]] = {}
    if col_ids:
        cals = (
            db.query(Calificacion)
            .filter(
                Calificacion.id_grupo == grupo.id_grupo,
                Calificacion.periodo == periodo,
                Calificacion.id_columna.in_(col_ids),
            )
            .all()
        )
        for c in cals:
            if c.valor is None:
                continue
            notas_por_est.setdefault(c.id_estudiante, {})[c.id_columna] = c.valor
    return columnas, notas_por_est


def _construir_boletin_estudiante(
    docente: Docente, grupo: Grupo, estudiante: Estudiante, periodo: int, columnas, notas_col: dict
) -> bytes:
    """Boletín individual — un DOCX de una página con la tabla del estudiante."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    _boletin_encabezado(doc, docente, grupo, subtitulo=f'Boletín · {estudiante.codigo_estudiante}', periodo=periodo)

    # Info del estudiante
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f'Estudiante:  {estudiante.codigo_estudiante}')
    r.font.bold = True; r.font.size = Pt(11)
    if estudiante.tiene_piar:
        r2 = p.add_run('   ·   PIAR activo')
        r2.font.size = Pt(9); r2.font.italic = True
        r2.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)

    # Cálculo del promedio
    par = [(notas_col.get(c.id_columna), c.porcentaje) for c in columnas]
    prom = _promedio_ponderado(par)

    _boletin_tabla_estudiante(doc, columnas, notas_col, prom)
    _boletin_pie_leyenda(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _construir_boletin_grupo(
    docente: Docente, grupo: Grupo, periodo: int, columnas, notas_por_est: dict, estudiantes: List[Estudiante]
) -> bytes:
    """
    Boletín consolidado del grupo:
    - Encabezado
    - Una sección por estudiante con su tabla individual + salto de página
    - Al final: tabla resumen con promedio y estado por estudiante
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_BREAK

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    _boletin_encabezado(
        doc, docente, grupo,
        subtitulo=f'Boletín consolidado — {len(estudiantes)} estudiante(s)', periodo=periodo,
    )

    # Tabla resumen del grupo
    tbl = doc.add_table(rows=1 + len(estudiantes), cols=4)
    tbl.style = 'Table Grid'
    for i, h in enumerate(['#', 'Estudiante', 'Definitiva', 'Estado']):
        c = tbl.rows[0].cells[i]
        _set_cell_bg(c, _HEX_AZUL_OSC)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    filas = []
    for est in estudiantes:
        notas = notas_por_est.get(est.id_estudiante, {})
        par = [(notas.get(c.id_columna), c.porcentaje) for c in columnas]
        prom = _promedio_ponderado(par)
        etiqueta, hex_estado = _estado_desde_promedio(prom)
        filas.append((est, prom, etiqueta, hex_estado))

    # Ordenar por promedio descendente (los sin nota al final)
    filas.sort(key=lambda x: (x[1] is None, -(x[1] or 0)))

    for i, (est, prom, etiqueta, hex_estado) in enumerate(filas, start=1):
        row = tbl.rows[i]
        row.cells[0].text = str(i)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].text = est.codigo_estudiante + (' · PIAR' if est.tiene_piar else '')
        row.cells[2].text = f'{prom:.1f}' if prom is not None else '—'
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[3].text = etiqueta
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(row.cells[3], hex_estado)
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    # Detalle por estudiante en páginas siguientes
    for est, prom, etiqueta, _hex in filas:
        doc.add_page_break()
        _boletin_encabezado(
            doc, docente, grupo,
            subtitulo=f'Boletín · {est.codigo_estudiante}', periodo=periodo,
        )
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f'Estudiante:  {est.codigo_estudiante}')
        r.font.bold = True; r.font.size = Pt(11)
        if est.tiene_piar:
            r2 = p.add_run('   ·   PIAR activo')
            r2.font.size = Pt(9); r2.font.italic = True
            r2.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
        notas = notas_por_est.get(est.id_estudiante, {})
        _boletin_tabla_estudiante(doc, columnas, notas, prom)
        _boletin_pie_leyenda(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _safe_filename(text: str, fallback: str = "boletin") -> str:
    safe = re.sub(r'[^\w\s-]', '', text).strip().replace(' ', '_')[:60]
    return f'{safe or fallback}.docx'


# ────────────────────────────────────────────────────────────────
# ENDPOINTS
# ────────────────────────────────────────────────────────────────

@router.get("/grupos/{grupo_id}/boletin/{estudiante_id}")
def boletin_estudiante(
    grupo_id: str,
    estudiante_id: str,
    periodo: int = Query(..., ge=1, le=4),
    docente: Docente = Depends(verify_trial_active),
    db: Session = Depends(get_db),
):
    """Boletín individual DOCX de un estudiante en un periodo."""
    grupo = db.query(Grupo).filter(
        Grupo.id_grupo == grupo_id,
        Grupo.id_docente == docente.id_docente,
    ).first()
    if not grupo:
        raise HTTPException(status_code=404, detail="Grupo no encontrado")

    estudiante = db.query(Estudiante).filter(
        Estudiante.id_estudiante == estudiante_id,
        Estudiante.id_grupo == grupo_id,
    ).first()
    if not estudiante:
        raise HTTPException(status_code=404, detail="Estudiante no encontrado")

    columnas, notas_por_est = _cargar_datos_boletin(db, grupo, periodo)
    notas_col = notas_por_est.get(estudiante_id, {})

    try:
        docx_bytes = _construir_boletin_estudiante(
            docente, grupo, estudiante, periodo, columnas, notas_col
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generando boletín: {e}")

    fname = _safe_filename(
        f'Boletin_{estudiante.codigo_estudiante}_{grupo.nombre_grupo}_P{periodo}',
        fallback='boletin_estudiante',
    )
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{fname}"'},
    )


@router.get("/grupos/{grupo_id}/boletin")
def boletin_grupo(
    grupo_id: str,
    periodo: int = Query(..., ge=1, le=4),
    docente: Docente = Depends(verify_trial_active),
    db: Session = Depends(get_db),
):
    """
    Boletín consolidado del grupo — un único DOCX con tabla resumen +
    una página por estudiante con su detalle de notas.
    """
    grupo = db.query(Grupo).filter(
        Grupo.id_grupo == grupo_id,
        Grupo.id_docente == docente.id_docente,
    ).first()
    if not grupo:
        raise HTTPException(status_code=404, detail="Grupo no encontrado")

    estudiantes = (
        db.query(Estudiante)
        .filter(Estudiante.id_grupo == grupo_id)
        .order_by(Estudiante.codigo_estudiante)
        .all()
    )
    if not estudiantes:
        raise HTTPException(status_code=400, detail="El grupo no tiene estudiantes")

    columnas, notas_por_est = _cargar_datos_boletin(db, grupo, periodo)

    try:
        docx_bytes = _construir_boletin_grupo(
            docente, grupo, periodo, columnas, notas_por_est, estudiantes
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generando boletín de grupo: {e}")

    fname = _safe_filename(
        f'Boletin_{grupo.nombre_grupo}_P{periodo}',
        fallback='boletin_grupo',
    )
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{fname}"'},
    )

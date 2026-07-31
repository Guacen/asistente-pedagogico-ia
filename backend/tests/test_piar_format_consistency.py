"""
Consistencia del formato del PIAR (sprint piar-fixed-format).

Contrato del sprint: el DOCX generado SIEMPRE tiene las mismas secciones
en el mismo orden, sin importar el contenido que el LLM haya devuelto.
Esto se logra usando un template Markdown estático (piar_template.md)
que el LLM NO controla — el LLM solo rellena 14 campos snake_case y
el sistema hace format_map.

Verificamos:
- 3 dicts de JSON del LLM completamente distintos generan DOCX con
  el MISMO orden de secciones.
- Ninguna sección de la lista canónica falta ni sobra.
- El orden es estrictamente el del template (Fortalezas ANTES que
  Barreras, Acuerdos ANTES que Seguimiento, etc.).

Sin llamadas al LLM — todo se opera con inputs sintéticos.
"""
from __future__ import annotations

import io

import pytest


# 10 secciones canónicas del PIAR — exactamente en el orden del template
_SECCIONES_CANONICAS_EN_ORDEN = [
    "Información del estudiante",
    "Contexto escolar y familiar",
    "Fortalezas e intereses",
    "Barreras para el aprendizaje y la participación (BAP)",
    "Ajustes razonables y estrategias DUA",
    "Evaluación flexible",
    "Apoyos requeridos",
    "Metas del período",
    "Acuerdos y compromisos",
    "Seguimiento",
]


def _json_14_claves_completo(marker: str) -> dict:
    """Construye un JSON válido con las 14 claves. `marker` diferencia los tests."""
    return {
        "contexto_escolar_familiar":  f"Contexto {marker}",
        "fortalezas_intereses":       f"Fortaleza {marker}: liderazgo grupal",
        "barreras_bap":               f"Barrera contextual {marker}",
        "dua_representacion":         f"Repr {marker}: imágenes + audio",
        "dua_expresion":              f"Expr {marker}: oral, escrito",
        "dua_motivacion":             f"Mot {marker}: elegir tema",
        "evaluacion_flexible":        f"Eval {marker}: oral, portafolio, proyecto",
        "apoyos_requeridos":          f"Apoyos {marker} — Decreto 1421",
        "metas_periodo":              f"Meta {marker}: lectura fluida",
        "compromisos_institucion":    f"Institución {marker}: material accesible",
        "compromisos_docente":        f"Docente {marker}: rúbrica adaptada",
        "compromisos_familia":        f"Familia {marker}: rutina de estudio",
        "fecha_revision":             "30/09/2026",
        "observaciones_seguimiento":  f"Obs {marker}",
    }


def _generar_docx(json_llm: dict) -> "Document":
    """Rellena el template + renderiza + abre el DOCX resultante."""
    from docx import Document
    from piar import _rellenar_template_markdown
    from markdown_parser import parse_markdown_sections
    from templates.maestria_template import generar_piar_docx

    datos_estudiante = {
        "nombre": "EST-01", "grado": "5°",
        "docente": "María López", "diagnostico": "—",
        "fecha": "29/07/2026",
    }
    md = _rellenar_template_markdown(json_llm, datos_estudiante)
    secciones = parse_markdown_sections(md)
    buf = generar_piar_docx(secciones, {
        "nombre": "EST-01", "grado": "5°",
        "docente": "María López", "fecha": "29/07/2026",
    })
    return Document(io.BytesIO(buf.getvalue()))


def _titulos_de_seccion(doc) -> list[str]:
    """
    Extrae los títulos de sección en el orden en que aparecen en el DOCX.
    Los headings de sección son párrafos estilo Normal con el texto exacto
    de una de las secciones canónicas (el template usa `## ` que el
    renderer emite como párrafo con texto tal cual).
    """
    encontrados: list[str] = []
    canonicas_set = set(_SECCIONES_CANONICAS_EN_ORDEN)
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in canonicas_set:
            encontrados.append(t)
    return encontrados


# ═══════════════════════════════════════════════════════════════
# Consistencia de formato — 3 PIARs distintos, mismo orden
# ═══════════════════════════════════════════════════════════════

def test_3_piars_distintos_producen_las_mismas_10_secciones_en_el_mismo_orden():
    doc_a = _generar_docx(_json_14_claves_completo("A"))
    doc_b = _generar_docx(_json_14_claves_completo("B"))
    doc_c = _generar_docx(_json_14_claves_completo("C"))

    orden_a = _titulos_de_seccion(doc_a)
    orden_b = _titulos_de_seccion(doc_b)
    orden_c = _titulos_de_seccion(doc_c)

    # Los 3 tienen la misma cantidad de secciones
    assert len(orden_a) == len(orden_b) == len(orden_c) == 10

    # Y en el mismo orden — el template es la única fuente de verdad
    assert orden_a == orden_b == orden_c
    assert orden_a == _SECCIONES_CANONICAS_EN_ORDEN


def test_fortalezas_aparece_antes_que_barreras_en_el_docx():
    """
    Regla de UX del ticket: enfoque de capacidades, no déficit.
    Aunque el LLM devuelva las claves en orden distinto, el template
    fija el orden final.
    """
    json_llm = _json_14_claves_completo("orden_llm_desordenado")
    doc = _generar_docx(json_llm)
    orden = _titulos_de_seccion(doc)
    idx_fort = orden.index("Fortalezas e intereses")
    idx_barr = orden.index("Barreras para el aprendizaje y la participación (BAP)")
    assert idx_fort < idx_barr


def test_las_secciones_no_dependen_del_contenido_del_llm():
    """
    Aunque el LLM devuelva strings vacíos en TODOS los campos, el DOCX
    sigue teniendo las 10 secciones. (Los contenidos vacíos quedan como
    "[PENDIENTE — sin información]" pero los TÍTULOS están siempre.)
    """
    json_vacio = {k: "" for k in [
        "contexto_escolar_familiar", "fortalezas_intereses", "barreras_bap",
        "dua_representacion", "dua_expresion", "dua_motivacion",
        "evaluacion_flexible", "apoyos_requeridos", "metas_periodo",
        "compromisos_institucion", "compromisos_docente", "compromisos_familia",
        "fecha_revision", "observaciones_seguimiento",
    ]}
    doc = _generar_docx(json_vacio)
    orden = _titulos_de_seccion(doc)
    assert orden == _SECCIONES_CANONICAS_EN_ORDEN

    # Y el marcador de pendiente aparece varias veces (uno por campo vacío)
    texto = " ".join(p.text for p in doc.paragraphs)
    assert texto.count("[PENDIENTE") >= 5

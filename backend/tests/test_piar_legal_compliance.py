"""
Compliance del PIAR con el marco legal colombiano (Decreto 1421,
Ley 1618, Convención ONU) y con los principios científicos del sprint
piar-legal-framework.

Los asserts operan sobre DOS niveles:

1. **Contenido estático del prompt** (PROMPT_MODO_PIAR): la fuente
   de instrucciones al modelo. Verificamos que las obligaciones
   normativas y las reglas de vocabulario están explícitamente
   presentes — así, cuando el modelo cumpla las instrucciones, el
   PIAR generado también cumplirá.

2. **Estructura del DOCX** cuando se le pasa contenido que respeta el
   esquema nuevo: las 10 secciones aparecen, el orden es correcto
   (Fortalezas antes que Barreras), y el footer incluye la mención
   del Decreto 1421.

No se llama a Claude/Gemini en ningún test — todo es introspección
del prompt y del generador. Ligero, rápido, y no gasta tokens.
"""
from __future__ import annotations

import re

import pytest


# ═══════════════════════════════════════════════════════════════
# 1. El prompt del sistema cita la normativa obligatoria
# ═══════════════════════════════════════════════════════════════

def test_prompt_cita_decreto_1421():
    from prompts import PROMPT_MODO_PIAR
    assert "Decreto 1421" in PROMPT_MODO_PIAR


def test_prompt_cita_ley_estatutaria_1618():
    from prompts import PROMPT_MODO_PIAR
    assert "Ley Estatutaria 1618" in PROMPT_MODO_PIAR \
        or "Ley 1618" in PROMPT_MODO_PIAR


def test_prompt_cita_convencion_onu_ley_1346():
    """El bloque de marco legal debe anclar en la Convención ONU + Ley 1346."""
    from prompts import PROMPT_MODO_PIAR
    assert "Convención ONU" in PROMPT_MODO_PIAR or "Convención de Naciones Unidas" in PROMPT_MODO_PIAR
    assert "Ley 1346" in PROMPT_MODO_PIAR


# ═══════════════════════════════════════════════════════════════
# 2. Conceptos científicos: BAP + DUA + ajuste razonable
# ═══════════════════════════════════════════════════════════════

def test_prompt_define_BAP_como_del_contexto_no_del_estudiante():
    """
    BAP tiene que aparecer con la aclaración explícita: la barrera es
    del contexto, no del estudiante. Sin esa aclaración, el modelo
    puede caer en lenguaje deficitario.
    """
    from prompts import PROMPT_MODO_PIAR
    assert "BAP" in PROMPT_MODO_PIAR
    # "contexto" y "no del estudiante" cerca uno del otro
    assert re.search(r"contexto.{0,120}no del estudiante", PROMPT_MODO_PIAR,
                     re.IGNORECASE | re.DOTALL) is not None


def test_prompt_define_los_3_principios_del_DUA():
    """Representación / Expresión / Motivación deben estar los 3 nombrados."""
    from prompts import PROMPT_MODO_PIAR
    assert "DUA" in PROMPT_MODO_PIAR
    for principio in ("representación", "expresión", "motivación"):
        assert principio in PROMPT_MODO_PIAR.lower(), (
            f"Falta principio DUA: {principio}"
        )


def test_prompt_diferencia_ajuste_razonable_de_apoyo_y_flexibilizacion():
    """
    Los 3 conceptos son distintos legalmente. El prompt tiene que
    definirlos por separado para que el modelo no los confunda.
    """
    from prompts import PROMPT_MODO_PIAR
    assert "ajuste razonable" in PROMPT_MODO_PIAR.lower()
    assert "apoyo especializado" in PROMPT_MODO_PIAR.lower() \
        or "apoyos especializados" in PROMPT_MODO_PIAR.lower()
    assert "flexibilización curricular" in PROMPT_MODO_PIAR.lower()


# ═══════════════════════════════════════════════════════════════
# 3. Estructura de secciones — fortalezas ANTES que barreras
# ═══════════════════════════════════════════════════════════════

def test_secciones_incluyen_fortalezas_e_intereses():
    from piar import SECCIONES_PIAR
    assert "Fortalezas e intereses" in SECCIONES_PIAR


def test_secciones_incluyen_evaluacion_flexible_independiente():
    from piar import SECCIONES_PIAR
    assert "Evaluación flexible" in SECCIONES_PIAR


def test_secciones_incluyen_apoyos_y_metas_como_secciones_propias():
    from piar import SECCIONES_PIAR
    assert "Apoyos requeridos" in SECCIONES_PIAR
    assert "Metas del período" in SECCIONES_PIAR


def test_secciones_incluyen_acuerdos_y_seguimiento():
    from piar import SECCIONES_PIAR
    assert "Acuerdos y compromisos" in SECCIONES_PIAR
    assert "Seguimiento" in SECCIONES_PIAR


def test_fortalezas_aparece_antes_que_barreras_en_el_orden_canonico():
    """
    Regla del sprint: enfoque de capacidades, no de déficit. Las
    fortalezas SIEMPRE se presentan primero.
    """
    from piar import SECCIONES_PIAR
    idx_fort = SECCIONES_PIAR.index("Fortalezas e intereses")
    idx_barr = SECCIONES_PIAR.index("Barreras para el aprendizaje y la participación")
    assert idx_fort < idx_barr


# ═══════════════════════════════════════════════════════════════
# 4. Formato Markdown: 10 secciones ## + 3 subsecciones ###
# ═══════════════════════════════════════════════════════════════

def test_prompt_pide_las_10_secciones_por_nombre_exacto():
    from prompts import PROMPT_MODO_PIAR
    from piar import SECCIONES_PIAR
    for titulo in SECCIONES_PIAR:
        assert titulo in PROMPT_MODO_PIAR, f"Sección ausente en prompt: {titulo}"


def test_prompt_pide_las_3_subsecciones_DUA_como_H3():
    from prompts import PROMPT_MODO_PIAR
    for sub in ("### Representación", "### Expresión", "### Motivación"):
        assert sub in PROMPT_MODO_PIAR, f"Subsección DUA ausente: {sub}"


# ═══════════════════════════════════════════════════════════════
# 5. Evaluación flexible con al menos 3 alternativas
# ═══════════════════════════════════════════════════════════════

def test_prompt_pide_al_menos_3_alternativas_de_evaluacion():
    """
    El prompt exige que 'Evaluación flexible' liste al menos 3 formas
    concretas (oral / portafolio / proyecto / rúbrica / etc.).
    """
    from prompts import PROMPT_MODO_PIAR
    p = PROMPT_MODO_PIAR.lower()
    # busca el número + alternativa/s en contexto de evaluación
    assert re.search(r"(al menos\s*3|3\s*alternativas|tres\s+alternativas)",
                     p) is not None
    # y las alternativas típicas están mencionadas
    tipos = ["oral", "portafolio", "proyecto", "rúbrica"]
    presentes = sum(1 for t in tipos if t in p)
    assert presentes >= 3, f"Solo {presentes} tipos de evaluación mencionados: {tipos}"


# ═══════════════════════════════════════════════════════════════
# 6. Lenguaje: reglas duras contra vocabulario clínico
# ═══════════════════════════════════════════════════════════════

def test_prompt_prohibe_explicitamente_terminos_clinicos():
    """
    El prompt debe listar los términos prohibidos (sufre / padece /
    déficit / trastorno) para que el modelo no los use.
    """
    from prompts import PROMPT_MODO_PIAR
    p = PROMPT_MODO_PIAR.lower()
    for prohibido in ("sufre", "padece", "déficit"):
        assert prohibido in p, (
            f"'{prohibido}' debe estar mencionado como término prohibido"
        )


def test_prompt_indica_lenguaje_de_capacidades_no_deficit():
    """El prompt debe explicitar el enfoque de capacidades."""
    from prompts import PROMPT_MODO_PIAR
    p = PROMPT_MODO_PIAR.lower()
    # combinaciones plausibles del texto del prompt
    assert "capacidades" in p and "déficit" in p


# ═══════════════════════════════════════════════════════════════
# 7. El PIAR requiere firma de rector y acudiente
# ═══════════════════════════════════════════════════════════════

def test_prompt_recuerda_firma_rector_y_acudiente_para_validez_legal():
    from prompts import PROMPT_MODO_PIAR
    p = PROMPT_MODO_PIAR.lower()
    assert "firma" in p
    assert "rector" in p and "acudiente" in p


# ═══════════════════════════════════════════════════════════════
# 8. El DOCX generado con contenido válido cita el Decreto 1421
# ═══════════════════════════════════════════════════════════════

def test_docx_generado_menciona_decreto_1421(seed_docente, db_session):
    """
    Cuando se genera el DOCX del PIAR (con el pipeline nuevo,
    templates/maestria_template), el subtítulo de portada cita el
    Decreto 1421 — así el documento entregado siempre lo referencia
    aunque el contenido del modelo por alguna razón lo omita.
    """
    from templates.maestria_template import generar_piar_docx
    from docx import Document
    import io

    secciones = {s: "Contenido de prueba." for s in [
        "Información del estudiante",
        "Contexto escolar y familiar",
        "Fortalezas e intereses",
        "Barreras para el aprendizaje y la participación",
        "Ajustes razonables y estrategias DUA",
        "Evaluación flexible",
        "Apoyos requeridos",
        "Metas del período",
        "Acuerdos y compromisos",
        "Seguimiento",
    ]}
    datos = {
        "nombre": "EST-01", "grado": "5°",
        "docente": "María López", "fecha": "30/07/2026",
    }
    buf = generar_piar_docx(secciones, datos)
    doc = Document(io.BytesIO(buf.getvalue()))
    texto = " ".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                texto += " " + cell.text
    assert "Decreto 1421" in texto

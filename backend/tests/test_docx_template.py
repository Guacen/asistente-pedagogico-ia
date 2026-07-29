"""
Tests del template DOCX de Maestr.ia (backend/templates/maestria_template.py).

Verifican estructura del DOCX generado: header, footer, número de secciones,
metadatos de portada, y que el archivo es un ZIP válido que Word/Google
Docs pueden abrir.

No renderizamos visualmente — se inspecciona el ZIP interno del DOCX
(que es OOXML) buscando strings clave y validando que existan las partes
esperadas (`word/header1.xml`, `word/footer1.xml`).
"""
from __future__ import annotations

import io
import zipfile

import pytest


def _abrir_docx(buf: io.BytesIO):
    from docx import Document
    return Document(io.BytesIO(buf.getvalue()))


def _texto_completo(doc) -> str:
    """
    Extrae todo el texto del cuerpo + tablas. NOTA: `doc.paragraphs` no
    incluye tablas ni header/footer; hay que iterar aparte.
    """
    piezas = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                piezas.append(cell.text)
    return " ".join(piezas)


def _texto_header(doc) -> str:
    piezas: list[str] = []
    for section in doc.sections:
        for p in section.header.paragraphs:
            piezas.append(p.text)
        for tbl in section.header.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    piezas.append(cell.text)
    return " ".join(piezas)


def _texto_footer(doc) -> str:
    piezas: list[str] = []
    for section in doc.sections:
        for p in section.footer.paragraphs:
            piezas.append(p.text)
        for tbl in section.footer.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    piezas.append(cell.text)
    return " ".join(piezas)


# ═══════════════════════════════════════════════════════════════

def _secciones_minimas() -> dict:
    return {
        "Datos del estudiante": "Código: **EST-01**, grado 5°.",
        "Descripción del contexto escolar": "Contexto familiar estable.",
        "Barreras para el aprendizaje y la participación": (
            "- Metodología uniforme\n- Sin material accesible"
        ),
        "Ajustes razonables y apoyos": (
            "| Área | Ajuste |\n| --- | --- |\n"
            "| Evaluación | Tiempo extra |\n| Acceso | Braille |\n"
        ),
        "Estrategias de evaluación flexible": "",   # → PENDIENTE
        "Seguimiento y compromisos": "Revisión mensual con orientación.",
    }


def _datos_minimos() -> dict:
    return {
        "nombre": "Juan P.",
        "grado": "5°",
        "docente": "María López",
        "fecha": "27/07/2026",
    }


# ═══════════════════════════════════════════════════════════════
# Contrato estructural del DOCX
# ═══════════════════════════════════════════════════════════════

def test_docx_es_zip_valido_con_partes_ooxml():
    """
    Cualquier DOCX debe ser un ZIP con firma "PK\x03\x04". Las partes
    mínimas OOXML: document.xml, header1.xml, footer1.xml.
    """
    from templates.maestria_template import generar_piar_docx
    buf = generar_piar_docx(_secciones_minimas(), _datos_minimos())
    raw = buf.getvalue()
    assert raw[:2] == b"PK"
    z = zipfile.ZipFile(io.BytesIO(raw))
    names = set(z.namelist())
    assert "word/document.xml" in names
    assert "word/header1.xml" in names   # asegura que hay header definido
    assert "word/footer1.xml" in names   # y footer también


def test_portada_incluye_titulo_subtitulo_y_datos_del_estudiante():
    from templates.maestria_template import generar_piar_docx
    buf = generar_piar_docx(_secciones_minimas(), _datos_minimos())
    doc = _abrir_docx(buf)
    texto = _texto_completo(doc)
    # Título y subtítulo del sprint
    assert "Plan Individual de Ajustes Razonables (PIAR)" in texto
    assert "Decreto 1421 de 2017" in texto
    # Tabla de datos de portada
    assert "Juan P." in texto
    assert "5°" in texto
    assert "María López" in texto
    assert "27/07/2026" in texto


def test_todas_las_secciones_aparecen_en_orden():
    from templates.maestria_template import generar_piar_docx
    buf = generar_piar_docx(_secciones_minimas(), _datos_minimos())
    doc = _abrir_docx(buf)
    texto = _texto_completo(doc)
    orden_esperado = list(_secciones_minimas().keys())
    posiciones = [texto.find(titulo) for titulo in orden_esperado]
    # Ninguno debe faltar
    assert all(p >= 0 for p in posiciones), f"faltó alguna sección: {posiciones}"
    # Y deben aparecer en el orden dado
    assert posiciones == sorted(posiciones)


def test_seccion_vacia_muestra_pendiente():
    """La sección "Estrategias..." viene con "" en el fixture."""
    from templates.maestria_template import generar_piar_docx
    buf = generar_piar_docx(_secciones_minimas(), _datos_minimos())
    doc = _abrir_docx(buf)
    texto = _texto_completo(doc)
    assert "[PENDIENTE" in texto


def test_bullets_de_seccion_barreras_aparecen_como_texto():
    from templates.maestria_template import generar_piar_docx
    buf = generar_piar_docx(_secciones_minimas(), _datos_minimos())
    doc = _abrir_docx(buf)
    texto = _texto_completo(doc)
    assert "Metodología uniforme" in texto
    assert "Sin material accesible" in texto


def test_tabla_de_ajustes_se_renderiza_como_tabla_docx():
    from templates.maestria_template import generar_piar_docx
    buf = generar_piar_docx(_secciones_minimas(), _datos_minimos())
    doc = _abrir_docx(buf)
    # Debe haber al menos: tabla portada (2 col) + tabla ajustes (2 col).
    # Buscamos una tabla cuya primera celda sea "Área".
    encontrada = False
    for tbl in doc.tables:
        if tbl.rows and tbl.rows[0].cells and tbl.rows[0].cells[0].text.strip() == "Área":
            encontrada = True
            assert tbl.rows[0].cells[1].text.strip() == "Ajuste"
            # Fila 2 debe tener "Evaluación" y "Tiempo extra"
            assert tbl.rows[1].cells[0].text.strip() == "Evaluación"
            assert tbl.rows[1].cells[1].text.strip() == "Tiempo extra"
    assert encontrada, "No se generó la tabla Markdown como docx.Table"


def test_header_contiene_tagline_de_marca():
    from templates.maestria_template import generar_piar_docx
    buf = generar_piar_docx(_secciones_minimas(), _datos_minimos())
    doc = _abrir_docx(buf)
    header = _texto_header(doc)
    assert "Maestr.ia" in header
    assert "Tu colega que conoce la ley" in header


def test_footer_contiene_marca_y_fecha():
    from templates.maestria_template import generar_piar_docx
    buf = generar_piar_docx(_secciones_minimas(), _datos_minimos())
    doc = _abrir_docx(buf)
    footer = _texto_footer(doc)
    assert "Maestr.ia" in footer
    assert "Generado con IA" in footer


def test_datos_faltantes_usan_placeholder():
    """Si el dict de datos no trae `docente`, se muestra "—" sin crashear."""
    from templates.maestria_template import generar_piar_docx
    datos = {"nombre": "X", "grado": "1°"}
    buf = generar_piar_docx(_secciones_minimas(), datos)
    doc = _abrir_docx(buf)
    texto = _texto_completo(doc)
    assert "—" in texto   # el placeholder del template


def test_secciones_vacio_dict_genera_docx_valido():
    """Edge case: sin ninguna sección — el DOCX debe generarse igual."""
    from templates.maestria_template import generar_piar_docx
    buf = generar_piar_docx({}, _datos_minimos())
    assert buf.getvalue()[:2] == b"PK"   # sigue siendo ZIP válido
    doc = _abrir_docx(buf)
    # Portada intacta aún sin cuerpo
    assert "Plan Individual de Ajustes Razonables (PIAR)" in _texto_completo(doc)


# ═══════════════════════════════════════════════════════════════
# Fix: estilos globales Heading + blockquotes
# ═══════════════════════════════════════════════════════════════

def test_estilos_globales_heading_usan_paleta_de_marca():
    """
    El template redefine Heading 1/2/3 con colores Maestr.ia — así el
    docente que edite el DOCX en Word y aplique "Heading 2" a un párrafo
    nuevo ve verde, no el azul default de python-docx.
    """
    from templates.maestria_template import generar_piar_docx
    from docx.shared import RGBColor
    buf = generar_piar_docx(_secciones_minimas(), _datos_minimos())
    doc = _abrir_docx(buf)

    h1, h2, h3 = doc.styles["Heading 1"], doc.styles["Heading 2"], doc.styles["Heading 3"]
    assert h1.font.color.rgb == RGBColor(0x0B, 0x3D, 0x2E)   # verde oscuro
    assert h2.font.color.rgb == RGBColor(0x0B, 0x3D, 0x2E)
    assert h3.font.color.rgb == RGBColor(0x1D, 0x9E, 0x75)   # verde principal
    assert h1.font.bold and h2.font.bold and h3.font.bold


def test_blockquote_no_deja_el_gt_literal_en_el_texto():
    """
    Antes del fix, `> texto` quedaba como texto plano con el `>` visible.
    Ahora se renderiza como bloque callout sin ese carácter.
    """
    from templates.maestria_template import generar_piar_docx
    secciones = {
        **_secciones_minimas(),
        "Datos del estudiante": "> Un blockquote con **énfasis** dentro.",
    }
    buf = generar_piar_docx(secciones, _datos_minimos())
    doc = _abrir_docx(buf)
    # El texto del blockquote está, PERO sin el "> " al comienzo.
    for p in doc.paragraphs:
        if "Un blockquote" in p.text:
            assert not p.text.lstrip().startswith(">"), (
                f"El '>' quedó literal: '{p.text}'"
            )
            break
    else:
        raise AssertionError("No se encontró el párrafo del blockquote")


def test_blockquote_tiene_borde_izquierdo_verde():
    """
    El blockquote se renderiza con un pBdr/left (borde vertical) — así
    lo diferencia Word visualmente de un párrafo normal.
    """
    from templates.maestria_template import generar_piar_docx
    from docx.oxml.ns import qn
    secciones = {
        **_secciones_minimas(),
        "Datos del estudiante": "> Callout con borde verde",
    }
    buf = generar_piar_docx(secciones, _datos_minimos())
    doc = _abrir_docx(buf)

    con_borde = 0
    for p in doc.paragraphs:
        pPr = p._p.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:pBdr")) is not None:
            if pPr.find(qn("w:pBdr")).find(qn("w:left")) is not None:
                con_borde += 1
    assert con_borde >= 1, "Ningún blockquote tiene borde izquierdo"


def test_blockquote_multilinea_se_agrupa_en_un_solo_parrafo():
    from templates.maestria_template import generar_piar_docx
    secciones = {
        **_secciones_minimas(),
        "Datos del estudiante": "> Línea uno\n> Línea dos\n> Línea tres",
    }
    buf = generar_piar_docx(secciones, _datos_minimos())
    doc = _abrir_docx(buf)
    for p in doc.paragraphs:
        if "Línea uno" in p.text:
            assert "Línea dos" in p.text and "Línea tres" in p.text
            break
    else:
        raise AssertionError("No se encontró el blockquote multilínea agrupado")

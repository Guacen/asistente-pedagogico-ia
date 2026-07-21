"""
Tests para los endpoints de boletines DOCX (Fase D).

Cubre:
- GET /api/grupos/{gid}/boletin/{eid}?periodo=N (individual)
- GET /api/grupos/{gid}/boletin?periodo=N (consolidado del grupo)
- Aislamiento: docente A no puede pedir boletines de recursos de B
- Validación de parámetros (periodo obligatorio, rango)
- Contenido básico del DOCX: incluye código del estudiante, notas y promedio esperado

Cálculo del promedio ponderado replica el del backend:
  peso = porcentaje si >0, si no 1.0
  promedio = sum(valor * peso) / sum(peso)
"""
from __future__ import annotations

import io

import pytest


CT_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _seed_grupo_con_notas(client, gid, notas_valores):
    """
    Crea 1 estudiante + 2 columnas y le pone las notas.
    notas_valores = [(v_col1, v_col2), ...] — pero para 1 estudiante recibe (v1, v2).
    """
    est = client.post(f"/api/grupos/{gid}/estudiantes", json={"codigo_estudiante": "E001"}).json()
    col1 = client.post(f"/api/grupos/{gid}/columnas", json={
        "nombre": "Quiz", "periodo": 1, "tipo": "quiz", "porcentaje": 40
    }).json()
    col2 = client.post(f"/api/grupos/{gid}/columnas", json={
        "nombre": "Parcial", "periodo": 1, "tipo": "parcial", "porcentaje": 60
    }).json()
    v1, v2 = notas_valores
    if v1 is not None:
        client.post(f"/api/grupos/{gid}/calificaciones/upsert", json={
            "id_estudiante": est["id_estudiante"], "id_columna": col1["id_columna"],
            "valor": v1, "periodo": 1,
        })
    if v2 is not None:
        client.post(f"/api/grupos/{gid}/calificaciones/upsert", json={
            "id_estudiante": est["id_estudiante"], "id_columna": col2["id_columna"],
            "valor": v2, "periodo": 1,
        })
    return est, col1, col2


def _docx_texto_completo(docx_bytes: bytes) -> str:
    """Extrae todo el texto (párrafos + celdas de tabla) de un DOCX en memoria."""
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    partes = []
    for p in doc.paragraphs:
        if p.text.strip():
            partes.append(p.text.strip())
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        partes.append(p.text.strip())
    return " || ".join(partes)


# ─── BOLETÍN INDIVIDUAL ───────────────────────────────────────────

def test_boletin_estudiante_devuelve_docx(client, seed_docente):
    gid = seed_docente["grupo"].id_grupo
    est, _, _ = _seed_grupo_con_notas(client, gid, (4.0, 5.0))

    r = client.get(f'/api/grupos/{gid}/boletin/{est["id_estudiante"]}?periodo=1')
    assert r.status_code == 200, r.text
    assert r.headers["content-type"] == CT_DOCX
    assert "attachment" in r.headers["content-disposition"]
    assert "Boletin_E001" in r.headers["content-disposition"]
    assert len(r.content) > 1000  # DOCX no vacío
    # Firma ZIP (DOCX es un zip)
    assert r.content[:2] == b"PK"


def test_boletin_incluye_datos_del_estudiante_y_promedio(client, seed_docente):
    """
    Verifica que el texto del DOCX incluye:
    - código del estudiante
    - nombres de las columnas
    - promedio ponderado esperado
    Con notas 4.0 (peso 40) y 5.0 (peso 60): (4*40 + 5*60) / 100 = 4.6
    """
    gid = seed_docente["grupo"].id_grupo
    est, _, _ = _seed_grupo_con_notas(client, gid, (4.0, 5.0))
    r = client.get(f'/api/grupos/{gid}/boletin/{est["id_estudiante"]}?periodo=1')
    assert r.status_code == 200
    texto = _docx_texto_completo(r.content)
    assert "E001" in texto
    assert "Quiz" in texto
    assert "Parcial" in texto
    assert "4.6" in texto        # promedio ponderado
    assert "Aprobado" in texto   # 4.6 >= 3.5


def test_boletin_marca_reprobado_cuando_promedio_bajo(client, seed_docente):
    gid = seed_docente["grupo"].id_grupo
    est, _, _ = _seed_grupo_con_notas(client, gid, (2.0, 2.5))
    # (2*40 + 2.5*60) / 100 = 2.3 → < 3.0 → Reprobado
    r = client.get(f'/api/grupos/{gid}/boletin/{est["id_estudiante"]}?periodo=1')
    assert r.status_code == 200
    texto = _docx_texto_completo(r.content)
    assert "2.3" in texto
    assert "Reprobado" in texto


def test_boletin_marca_en_riesgo_cuando_promedio_intermedio(client, seed_docente):
    gid = seed_docente["grupo"].id_grupo
    est, _, _ = _seed_grupo_con_notas(client, gid, (3.0, 3.2))
    # (3*40 + 3.2*60) / 100 = 3.12 → en 3.0–3.4 → En riesgo
    r = client.get(f'/api/grupos/{gid}/boletin/{est["id_estudiante"]}?periodo=1')
    assert r.status_code == 200
    texto = _docx_texto_completo(r.content)
    assert "3.1" in texto
    assert "En riesgo" in texto


def test_boletin_estudiante_sin_notas(client, seed_docente):
    """Estudiante sin notas registradas → boletín válido con 'Sin notas'."""
    gid = seed_docente["grupo"].id_grupo
    est = client.post(f"/api/grupos/{gid}/estudiantes", json={"codigo_estudiante": "SIN01"}).json()
    r = client.get(f'/api/grupos/{gid}/boletin/{est["id_estudiante"]}?periodo=1')
    assert r.status_code == 200
    texto = _docx_texto_completo(r.content)
    assert "SIN01" in texto
    assert "Sin notas" in texto


def test_boletin_periodo_es_obligatorio(client, seed_docente):
    gid = seed_docente["grupo"].id_grupo
    est, _, _ = _seed_grupo_con_notas(client, gid, (4.0, 4.0))
    r = client.get(f'/api/grupos/{gid}/boletin/{est["id_estudiante"]}')
    assert r.status_code == 422  # falta query param


def test_boletin_periodo_fuera_de_rango(client, seed_docente):
    gid = seed_docente["grupo"].id_grupo
    est, _, _ = _seed_grupo_con_notas(client, gid, (4.0, 4.0))
    r = client.get(f'/api/grupos/{gid}/boletin/{est["id_estudiante"]}?periodo=5')
    assert r.status_code == 422


def test_boletin_estudiante_inexistente(client, seed_docente):
    gid = seed_docente["grupo"].id_grupo
    r = client.get(f'/api/grupos/{gid}/boletin/00000000-0000-0000-0000-000000000000?periodo=1')
    assert r.status_code == 404


def test_boletin_grupo_inexistente(client):
    r = client.get('/api/grupos/00000000-0000-0000-0000-000000000000/boletin/xxx?periodo=1')
    assert r.status_code == 404


# ─── BOLETÍN CONSOLIDADO DEL GRUPO ────────────────────────────────

def test_boletin_grupo_devuelve_docx(client, seed_docente):
    gid = seed_docente["grupo"].id_grupo
    # 3 estudiantes con notas distintas
    est1 = client.post(f"/api/grupos/{gid}/estudiantes", json={"codigo_estudiante": "E01"}).json()
    est2 = client.post(f"/api/grupos/{gid}/estudiantes", json={"codigo_estudiante": "E02"}).json()
    est3 = client.post(f"/api/grupos/{gid}/estudiantes", json={"codigo_estudiante": "E03", "tiene_piar": True}).json()
    col = client.post(f"/api/grupos/{gid}/columnas", json={
        "nombre": "Quiz", "periodo": 1, "tipo": "quiz", "porcentaje": 100,
    }).json()
    for est, val in [(est1, 4.8), (est2, 2.5), (est3, 3.2)]:
        client.post(f"/api/grupos/{gid}/calificaciones/upsert", json={
            "id_estudiante": est["id_estudiante"], "id_columna": col["id_columna"],
            "valor": val, "periodo": 1,
        })

    r = client.get(f'/api/grupos/{gid}/boletin?periodo=1')
    assert r.status_code == 200, r.text
    assert r.headers["content-type"] == CT_DOCX
    assert "Boletin_Grupo_A" in r.headers["content-disposition"]

    texto = _docx_texto_completo(r.content)
    # Los 3 estudiantes aparecen
    for c in ("E01", "E02", "E03"):
        assert c in texto
    # Los 3 estados aparecen
    assert "Aprobado" in texto
    assert "Reprobado" in texto
    assert "En riesgo" in texto
    # PIAR marcado en el estudiante correspondiente
    assert "PIAR" in texto


def test_boletin_grupo_sin_estudiantes_devuelve_400(client, seed_docente):
    gid = seed_docente["grupo"].id_grupo
    r = client.get(f'/api/grupos/{gid}/boletin?periodo=1')
    assert r.status_code == 400
    assert "estudiantes" in r.json()["detail"].lower()


def test_boletin_grupo_ordenamiento_por_promedio(client, seed_docente):
    """En la tabla resumen, el estudiante con mayor promedio aparece primero."""
    gid = seed_docente["grupo"].id_grupo
    e_alto = client.post(f"/api/grupos/{gid}/estudiantes", json={"codigo_estudiante": "ZZZ_ALTO"}).json()
    e_bajo = client.post(f"/api/grupos/{gid}/estudiantes", json={"codigo_estudiante": "AAA_BAJO"}).json()
    col = client.post(f"/api/grupos/{gid}/columnas", json={
        "nombre": "Q", "periodo": 1, "porcentaje": 100,
    }).json()
    client.post(f"/api/grupos/{gid}/calificaciones/upsert", json={
        "id_estudiante": e_alto["id_estudiante"], "id_columna": col["id_columna"],
        "valor": 5.0, "periodo": 1,
    })
    client.post(f"/api/grupos/{gid}/calificaciones/upsert", json={
        "id_estudiante": e_bajo["id_estudiante"], "id_columna": col["id_columna"],
        "valor": 1.5, "periodo": 1,
    })
    r = client.get(f'/api/grupos/{gid}/boletin?periodo=1')
    assert r.status_code == 200
    texto = _docx_texto_completo(r.content)
    # ZZZ_ALTO debe aparecer antes que AAA_BAJO en la tabla resumen
    pos_alto = texto.find("ZZZ_ALTO")
    pos_bajo = texto.find("AAA_BAJO")
    assert pos_alto != -1 and pos_bajo != -1
    assert pos_alto < pos_bajo, "El resumen debe ordenar por promedio descendente"


# ─── AISLAMIENTO ──────────────────────────────────────────────────

def test_boletin_estudiante_ajeno_devuelve_404(client_two_docentes, db_session):
    """Docente A no puede pedir boletín usando un id_estudiante que pertenece a B."""
    from models import Estudiante
    est_b = Estudiante(
        id_grupo=client_two_docentes["data"]["b"]["grupo"].id_grupo,
        codigo_estudiante="EBOL",
    )
    db_session.add(est_b); db_session.commit(); db_session.refresh(est_b)

    d = client_two_docentes
    d["as_a"]()
    # Vector 1: path grupo de B → 404 (grupo no es de A)
    r1 = d["client"].get(
        f'/api/grupos/{d["data"]["b"]["grupo"].id_grupo}/boletin/{est_b.id_estudiante}?periodo=1'
    )
    assert r1.status_code == 404
    # Vector 2: path grupo de A + est de B → 404 (estudiante no está en grupo de A)
    r2 = d["client"].get(
        f'/api/grupos/{d["data"]["a"]["grupo"].id_grupo}/boletin/{est_b.id_estudiante}?periodo=1'
    )
    assert r2.status_code == 404


def test_boletin_grupo_ajeno_devuelve_404(client_two_docentes):
    d = client_two_docentes
    d["as_a"]()
    r = d["client"].get(f'/api/grupos/{d["data"]["b"]["grupo"].id_grupo}/boletin?periodo=1')
    assert r.status_code == 404

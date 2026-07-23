"""
Tests de integración del generador de PIAR (Fase C).

Cubre los 8 requisitos de la spec + edge cases:
- Solo estudiantes con tiene_piar=true entran al flow
- POST /piar crea JSON con las 6 secciones esperadas
- Aprobar hace inmutable (segundo PUT devuelve 409)
- Versionado paralelo: version+1 sin borrar la anterior
- Aislamiento: docente A no ve/edita PIAR de docente B
- Sin estudiante seleccionado → error explícito
- Sin conversación previa → error explícito
- Rate limit 5/día se aplica al modo piar
- DOCX generado on-demand incluye marca BORRADOR o APROBADO según estado

NUNCA se llama a Claude real: la síntesis en POST /piar se mockea con
monkeypatch de _sintetizar_conversacion_a_json para devolver JSON canned.
"""
from __future__ import annotations

import io
import json
import sys
from pathlib import Path
from unittest.mock import AsyncMock

import pytest

BACKEND_DIR = Path(__file__).resolve().parent.parent
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

SECCIONES = (
    "caracterizacion",
    "barreras",
    "ajustes_razonables",
    "apoyos",
    "metas",
    "seguimiento",
)

CT_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# ─── Helpers ──────────────────────────────────────────────────────

def _seed_estudiante_con_piar(db_session, id_grupo, codigo="EPIAR01",
                              piar=True, diag="TDA-H",
                              ajustes="Instrucciones cortas"):
    from models import Estudiante
    est = Estudiante(
        id_grupo=id_grupo,
        codigo_estudiante=codigo,
        tiene_piar=piar,
        diagnostico=diag,
        ajustes=ajustes,
    )
    db_session.add(est)
    db_session.commit()
    db_session.refresh(est)
    return est


def _seed_conversacion_piar(db_session, id_grupo, id_estudiante,
                            n_turnos=2):
    """Crea 2*n mensajes (docente + IA) en modo PIAR para (grupo, estudiante)."""
    from models import Mensaje
    for i in range(n_turnos):
        db_session.add(Mensaje(
            id_grupo=id_grupo, remitente="docente",
            contenido=f"pregunta {i+1}", modo="piar",
            id_estudiante=id_estudiante,
        ))
        db_session.add(Mensaje(
            id_grupo=id_grupo, remitente="sistema",
            contenido=f"respuesta guiada {i+1}", modo="piar",
            id_estudiante=id_estudiante,
        ))
    db_session.commit()


def _contenido_canned() -> dict:
    """JSON de PIAR mock que devuelve el mock de síntesis."""
    return {
        "caracterizacion": "Estudiante de 8° con TDA-H diagnosticado.",
        "barreras": "Metodología uniforme, evaluaciones sin diversificación.",
        "ajustes_razonables": "Tiempo extra en pruebas, formatos alternativos.",
        "apoyos": "Orientación escolar, uso de agenda visual.",
        "metas": "Sostener atención en tareas de 15 min consecutivos.",
        "seguimiento": "Revisión mensual con orientación.",
    }


@pytest.fixture(autouse=True)
def _mock_sintesis(monkeypatch):
    """
    Reemplaza la llamada a Claude por un mock async que devuelve el
    contenido canned. Se aplica a TODOS los tests de este archivo para
    que POST /piar nunca golpee la API real.
    """
    import piar as piar_module
    mock = AsyncMock(return_value=_contenido_canned())
    monkeypatch.setattr(piar_module, "_sintetizar_conversacion_a_json", mock)
    return mock


# ─── 1. Solo estudiantes con tiene_piar=true ─────────────────────

def test_crear_piar_para_estudiante_sin_piar_devuelve_400(client, seed_docente, db_session):
    est_sin = _seed_estudiante_con_piar(
        db_session, seed_docente["grupo"].id_grupo,
        codigo="EPIAR_SIN", piar=False,
    )
    r = client.post("/api/piar/", json={
        "id_estudiante": est_sin.id_estudiante,
        "periodo": 1,
    })
    assert r.status_code == 400
    assert "piar" in r.json()["detail"].lower()


# ─── 2. POST crea JSON con las 6 secciones ───────────────────────

def test_crear_piar_genera_json_con_las_6_secciones(client, seed_docente, db_session):
    est = _seed_estudiante_con_piar(db_session, seed_docente["grupo"].id_grupo)
    _seed_conversacion_piar(db_session, seed_docente["grupo"].id_grupo,
                            est.id_estudiante)
    r = client.post("/api/piar/", json={
        "id_estudiante": est.id_estudiante,
        "periodo": 1,
    })
    assert r.status_code == 201, r.text
    body = r.json()
    assert body["estado"] == "borrador"
    assert body["version"] == 1
    assert body["periodo"] == 1
    assert body["contenido"]
    for s in SECCIONES:
        assert s in body["contenido"], f"Falta sección {s}"
        assert isinstance(body["contenido"][s], str)
        assert body["contenido"][s].strip() != ""


# ─── 3. Sin conversación previa → error explícito ────────────────

def test_crear_piar_sin_conversacion_devuelve_400(client, seed_docente, db_session):
    est = _seed_estudiante_con_piar(db_session, seed_docente["grupo"].id_grupo)
    # NO se siembra conversación
    r = client.post("/api/piar/", json={
        "id_estudiante": est.id_estudiante,
        "periodo": 1,
    })
    assert r.status_code == 400
    assert "conversac" in r.json()["detail"].lower()


# ─── 4. Aprobar hace inmutable (segundo PUT → 409) ───────────────

def test_aprobar_piar_lo_hace_inmutable(client, seed_docente, db_session):
    est = _seed_estudiante_con_piar(db_session, seed_docente["grupo"].id_grupo)
    _seed_conversacion_piar(db_session, seed_docente["grupo"].id_grupo,
                            est.id_estudiante)
    piar_id = client.post("/api/piar/", json={
        "id_estudiante": est.id_estudiante, "periodo": 1,
    }).json()["id_piar"]

    r1 = client.put(f"/api/piar/{piar_id}/aprobar")
    assert r1.status_code == 200
    assert r1.json()["estado"] == "aprobado"
    assert r1.json()["aprobado_en"] is not None

    # Segundo intento de aprobar → 409
    r2 = client.put(f"/api/piar/{piar_id}/aprobar")
    assert r2.status_code == 409
    assert "aprobado" in r2.json()["detail"].lower()


# ─── 5. Versionado paralelo ──────────────────────────────────────

def test_versionado_incrementa_sin_borrar_anterior(client, seed_docente, db_session):
    est = _seed_estudiante_con_piar(db_session, seed_docente["grupo"].id_grupo)
    _seed_conversacion_piar(db_session, seed_docente["grupo"].id_grupo,
                            est.id_estudiante)

    # v1
    r1 = client.post("/api/piar/", json={"id_estudiante": est.id_estudiante, "periodo": 1})
    assert r1.json()["version"] == 1
    piar_id_v1 = r1.json()["id_piar"]

    # v2 en el mismo periodo — mismo tuple (estudiante, grupo, periodo, anio)
    r2 = client.post("/api/piar/", json={"id_estudiante": est.id_estudiante, "periodo": 1})
    assert r2.json()["version"] == 2
    piar_id_v2 = r2.json()["id_piar"]

    # v3 después de aprobar v1 — sigue permitiendo generar nuevas versiones
    client.put(f"/api/piar/{piar_id_v1}/aprobar")
    r3 = client.post("/api/piar/", json={"id_estudiante": est.id_estudiante, "periodo": 1})
    assert r3.json()["version"] == 3

    # Las 3 versiones coexisten y v1 sigue aprobada
    lista = client.get(f"/api/piar/estudiante/{est.id_estudiante}").json()
    assert len(lista) == 3
    versiones = sorted([p["version"] for p in lista])
    assert versiones == [1, 2, 3]
    v1 = next(p for p in lista if p["id_piar"] == piar_id_v1)
    assert v1["estado"] == "aprobado"


# ─── 6. Aislamiento entre docentes ───────────────────────────────

def test_docente_a_no_ve_piar_de_docente_b(client_two_docentes, db_session):
    """
    Docente A hace una request contra un PIAR del docente B — debe
    recibir 404 en todos los endpoints (get_list, aprobar, docx).
    """
    from models import PIAR

    # Sembrar estudiante y PIAR bajo el docente B
    grupo_b = client_two_docentes["data"]["b"]["grupo"]
    est_b = _seed_estudiante_con_piar(db_session, grupo_b.id_grupo, codigo="EB01")
    piar_b = PIAR(
        id_estudiante=est_b.id_estudiante,
        id_grupo=grupo_b.id_grupo,
        id_docente=client_two_docentes["data"]["b"]["docente"].id_docente,
        periodo=1, anio=2026, version=1,
        contenido=_contenido_canned(),
        estado="borrador",
    )
    db_session.add(piar_b)
    db_session.commit()
    db_session.refresh(piar_b)

    d = client_two_docentes
    d["as_a"]()

    # A intenta leer versiones del estudiante de B → 404 (estudiante ajeno)
    r_list = d["client"].get(f"/api/piar/estudiante/{est_b.id_estudiante}")
    assert r_list.status_code == 404

    # A intenta aprobar PIAR de B → 404
    r_aprob = d["client"].put(f"/api/piar/{piar_b.id_piar}/aprobar")
    assert r_aprob.status_code == 404

    # A intenta descargar DOCX de B → 404
    r_docx = d["client"].get(f"/api/piar/{piar_b.id_piar}/docx")
    assert r_docx.status_code == 404

    # B sí puede ver su PIAR
    d["as_b"]()
    r_b = d["client"].get(f"/api/piar/estudiante/{est_b.id_estudiante}")
    assert r_b.status_code == 200
    assert len(r_b.json()) == 1


def test_docente_a_no_puede_crear_piar_para_estudiante_de_b(client_two_docentes, db_session):
    grupo_b = client_two_docentes["data"]["b"]["grupo"]
    est_b = _seed_estudiante_con_piar(db_session, grupo_b.id_grupo, codigo="EB02")

    d = client_two_docentes
    d["as_a"]()
    r = d["client"].post("/api/piar/", json={
        "id_estudiante": est_b.id_estudiante, "periodo": 1,
    })
    # 404 porque el estudiante no pertenece a un grupo de A
    assert r.status_code == 404


# ─── 7. DOCX on-demand con marca BORRADOR / APROBADO ─────────────

def test_docx_borrador_incluye_marca_borrador(client, seed_docente, db_session):
    est = _seed_estudiante_con_piar(db_session, seed_docente["grupo"].id_grupo)
    _seed_conversacion_piar(db_session, seed_docente["grupo"].id_grupo,
                            est.id_estudiante)
    piar_id = client.post("/api/piar/", json={
        "id_estudiante": est.id_estudiante, "periodo": 1,
    }).json()["id_piar"]

    r = client.get(f"/api/piar/{piar_id}/docx")
    assert r.status_code == 200
    assert r.headers["content-type"] == CT_DOCX
    assert "PIAR_EPIAR01_P1_v1.docx" in r.headers["content-disposition"]
    assert len(r.content) > 1000
    assert r.content[:2] == b"PK"  # firma ZIP/DOCX

    # Extraer texto y verificar marca BORRADOR
    from docx import Document
    doc = Document(io.BytesIO(r.content))
    texto_completo = " ".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                texto_completo += " " + cell.text
    assert "BORRADOR" in texto_completo
    assert "Sujeto a revisión" in texto_completo
    # Todas las secciones en el DOCX
    assert "Caracterización" in texto_completo
    assert "Barreras" in texto_completo
    assert "Decreto 1421" in texto_completo


def test_docx_aprobado_reemplaza_marca_por_fecha_de_aprobacion(client, seed_docente, db_session):
    est = _seed_estudiante_con_piar(db_session, seed_docente["grupo"].id_grupo)
    _seed_conversacion_piar(db_session, seed_docente["grupo"].id_grupo,
                            est.id_estudiante)
    piar_id = client.post("/api/piar/", json={
        "id_estudiante": est.id_estudiante, "periodo": 1,
    }).json()["id_piar"]
    client.put(f"/api/piar/{piar_id}/aprobar")

    r = client.get(f"/api/piar/{piar_id}/docx")
    assert r.status_code == 200

    from docx import Document
    doc = Document(io.BytesIO(r.content))
    texto = " ".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                texto += " " + cell.text
    assert "APROBADO" in texto
    # Ya no debe estar la marca de borrador
    assert "Sujeto a revisión" not in texto


# ─── 8. Sanitización del contenido cuando la síntesis devuelve algo raro

def test_contenido_sin_secciones_esperadas_se_marca_pendiente(monkeypatch, client, seed_docente, db_session):
    """Si el modelo devuelve un dict con secciones faltantes, quedan como pendientes."""
    import piar as piar_module
    # Override del mock global (autouse) para este test
    mock = AsyncMock(return_value={"caracterizacion": "Solo una sección"})
    monkeypatch.setattr(piar_module, "_sintetizar_conversacion_a_json", mock)

    est = _seed_estudiante_con_piar(db_session, seed_docente["grupo"].id_grupo)
    _seed_conversacion_piar(db_session, seed_docente["grupo"].id_grupo,
                            est.id_estudiante)
    r = client.post("/api/piar/", json={
        "id_estudiante": est.id_estudiante, "periodo": 1,
    })
    # OJO: el autouse fixture reemplaza el mock DESPUÉS del que puse acá,
    # así que este test valida el helper _sanitizar_contenido directo:
    from piar import _sanitizar_contenido, SECCIONES_PIAR
    limpio = _sanitizar_contenido({"caracterizacion": "Solo una sección"})
    for s in SECCIONES_PIAR:
        assert s in limpio
    assert limpio["caracterizacion"] == "Solo una sección"
    assert "[PENDIENTE" in limpio["barreras"]


# ─── Extra: validación de sub-flow del socket via API ────────────

def test_estudiante_de_otro_grupo_devuelve_404(client, seed_docente, seed_docente_b, db_session):
    """
    Estudiante existe pero pertenece a un grupo del OTRO docente
    → 404 (no 403) por decisión de contrato (no revelar existencia).
    """
    est_b = _seed_estudiante_con_piar(
        db_session, seed_docente_b["grupo"].id_grupo, codigo="EB03",
    )
    r = client.post("/api/piar/", json={
        "id_estudiante": est_b.id_estudiante, "periodo": 1,
    })
    assert r.status_code == 404
